# --- ADD: Module docstring ---
"""Module for processing Word documents (.docx)."""
# --- END ADD ---
import logging
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
import os
import zipfile
from lxml import etree
from .exceptions import (
    DocumentNotFoundError, InvalidDocumentError, DocumentLoadingError,
    DocumentSavingError, XMLParsingError, DocumentNotLoadedError
)

# --- ADD: Configure logging ---
logging.basicConfig(level=logging.INFO, format='[%(asctime)s] [%(levelname)s] [%(name)s.%(funcName)s] - %(message)s')
logger = logging.getLogger(__name__)
# --- END ADD ---

class WordProcessor:
    """
    Processes Word documents (.docx) and extracts review markings.
    Allows customization of output formatting for revisions and comments.
    """
    def __init__(self, filepath,
                 added_tag_start: str = "\\added{", added_tag_end: str = "}",
                 deleted_tag_start: str = "\\deleted{", deleted_tag_end: str = "}",
                 comment_tag_start: str = "\\comment{", comment_tag_end: str = "}",
                 merge_revisions: bool = True):
        """
        Initializes the WordProcessor with a document file path and custom tags.

        Args:
            filepath: The path to the .docx file.
            added_tag_start: The starting tag for added text (default: "\\added{").
            added_tag_end: The ending tag for added text (default: "}").
            deleted_tag_start: The starting tag for deleted text (default: "\\deleted{").
            deleted_tag_end: The ending tag for deleted text (default: "}").
            comment_tag_start: The starting tag for comments (default: "\\comment{").
            comment_tag_end: The ending tag for comments (default: "}").
            merge_revisions: Whether to merge revisions into a single string (default: True).
        """
        self._filepath = filepath
        self._document = None
        self._added_tag_start = added_tag_start
        self._added_tag_end = added_tag_end
        self._deleted_tag_start = deleted_tag_start
        self._deleted_tag_end = deleted_tag_end
        self._comment_tag_start = comment_tag_start
        self._comment_tag_end = comment_tag_end
        self._merge_revisions = merge_revisions
        logger.info(f"Initialized WordProcessor with file: {filepath}")

    def load_document(self):
        # --- ADD: Method docstring ---
        """
        Loads the Word document from the specified filepath.

        Raises:
            DocumentNotFoundError: If the document file does not exist.
            InvalidDocumentError: If the document is not a valid .docx file or is corrupt.
            DocumentLoadingError: For any other unexpected errors during loading.
        """
        # --- END ADD ---
        if not os.path.exists(self._filepath):
            logger.error(f"文件未找到: {self._filepath}")
            raise DocumentNotFoundError(self._filepath)
        try:
            self._document = Document(self._filepath)
            logger.info(f"成功加载文档: {self._filepath}")
        except PackageNotFoundError as e:
            logger.error(f"加载文档 {self._filepath} 错误: 不是有效的 .docx 文件或包已损坏。异常: {e}", exc_info=True)
            raise InvalidDocumentError(self._filepath, original_exception=e) from e
        except Exception as e:
            logger.error(f"加载文档 {self._filepath} 时发生意外错误。异常: {e}", exc_info=True)
            raise DocumentLoadingError(self._filepath, original_exception=e) from e

    def read_paragraphs(self):
        # --- ADD: Method docstring ---
        """
        Reads all paragraphs from the loaded document.

        Yields:
            str: The text content of each paragraph.

        Raises:
            DocumentNotLoadedError: If the document has not been loaded.
        """
        # --- END ADD ---
        if self._document is None:
            logger.error("文档未加载。请先调用 load_document()。")
            raise DocumentNotLoadedError()

        logger.info("正在从文档中读取段落。")
        paragraph_count = 0
        for paragraph in self._document.paragraphs:
            logger.debug(f"读取段落 (前50个字符): '{paragraph.text[:50]}'")
            yield paragraph.text
            paragraph_count += 1
        logger.info(f"完成读取 {paragraph_count} 个段落。")

    def read_comments(self):
        # --- ADD: Method docstring ---
        """
        Reads all comments from the loaded document.

        Yields:
            dict: A dictionary representing a comment with 'id', 'author', 'date', and 'text'.

        Raises:
            DocumentNotLoadedError: If the document has not been loaded.
            XMLParsingError: If there's an error parsing the comments XML.
        """
        # --- END ADD ---
        if self._document is None:
            logger.error("文档未加载。请先调用 load_document()。")
            raise DocumentNotLoadedError()

        logger.info("正在从文档中读取评论。")
        comment_count = 0
        comments_tree = self._get_comments_xml_tree()
        if comments_tree is None:
            logger.warning("无法获取评论 XML 树。无法读取评论。")
            return # _get_comments_xml_tree already logs/raises if there's a critical error

        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        # Find all <w:comment> elements
        for comment_element in comments_tree.xpath('//w:comment', namespaces=nsmap):
            comment_id = comment_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
            author = comment_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author')
            date = comment_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date')

            # Extract text from within the comment element
            # Comments can contain multiple <w:p> (paragraph) elements, each with <w:r> (run) and <w:t> (text)
            comment_text_parts = []
            for p_element in comment_element.xpath('.//w:p', namespaces=nsmap):
                for t_element in p_element.xpath('.//w:t', namespaces=nsmap):
                    comment_text_parts.append(t_element.text)
            comment_text = "".join(comment_text_parts)

            logger.debug(f"Read comment: id='{comment_id}', author='{author}', date='{date}', text='{comment_text[:50]}'")
            yield {
                "id": comment_id,
                "author": author,
                "date": date,
                "text": comment_text
            }
            comment_count += 1
        logger.info(f"Finished reading {comment_count} comments.")

    def _get_comments_xml_tree(self):
        """
        Extracts and parses the word/comments.xml from the .docx file.

        Returns:
            lxml.etree._ElementTree: The parsed XML tree of comments.xml, or None if comments.xml is not found.

        Raises:
            XMLParsingError: If there's an error extracting or parsing comments.xml.
        """
        try:
            with zipfile.ZipFile(self._filepath, 'r') as docx_zip:
                if 'word/comments.xml' in docx_zip.namelist():
                    with docx_zip.open('word/comments.xml') as comments_xml_file:
                        return etree.parse(comments_xml_file)
                else:
                    logger.info("文档中未找到 'word/comments.xml'。没有评论可读。")
                    return None
        except zipfile.BadZipFile as e:
            logger.error(f"提取 comments.xml 错误: {self._filepath} 不是有效的 zip 文件（.docx 已损坏）。异常: {e}", exc_info=True)
            raise XMLParsingError("word/comments.xml", original_exception=e) from e
        except Exception as e:
            logger.error(f"提取 comments.xml 时发生意外错误: {e}", exc_info=True)
            raise XMLParsingError("word/comments.xml", original_exception=e) from e

    def _get_document_xml_tree(self):
        """
        Extracts and parses the word/document.xml from the .docx file.

        Returns:
            lxml.etree._ElementTree: The parsed XML tree of document.xml.

        Raises:
            XMLParsingError: If 'word/document.xml' is not found or if there's an error extracting or parsing it.
        """
        try:
            with zipfile.ZipFile(self._filepath, 'r') as docx_zip:
                if 'word/document.xml' in docx_zip.namelist():
                    with docx_zip.open('word/document.xml') as document_xml_file:
                        return etree.parse(document_xml_file)
                else:
                    logger.error("文档中未找到 'word/document.xml'。")
                    raise XMLParsingError("word/document.xml", message="Document XML part not found")
        except zipfile.BadZipFile as e:
            logger.error(f"提取 document.xml 错误: {self._filepath} 不是有效的 zip 文件（.docx 已损坏）。异常: {e}", exc_info=True)
            raise XMLParsingError("word/document.xml", original_exception=e) from e
        except Exception as e:
            logger.error(f"提取 document.xml 时发生意外错误: {e}", exc_info=True)
            raise XMLParsingError("word/document.xml", original_exception=e) from e

    def get_document_xml_content(self):
        """
        Extracts and returns the raw XML content of word/document.xml.

        Returns:
            str: The XML content as a string.

        Raises:
            XMLParsingError: If 'word/document.xml' is not found or if there's an error extracting or parsing it.
        """
        try:
            with zipfile.ZipFile(self._filepath, 'r') as docx_zip:
                if 'word/document.xml' in docx_zip.namelist():
                    with docx_zip.open('word/document.xml') as document_xml_file:
                        return document_xml_file.read().decode('utf-8')
                else:
                    logger.error("文档中未找到 'word/document.xml'。")
                    raise XMLParsingError("word/document.xml", message="Document XML part not found")
        except zipfile.BadZipFile as e:
            logger.error(f"提取 document.xml 错误: {self._filepath} 不是有效的 zip 文件（.docx 已损坏）。异常: {e}", exc_info=True)
            raise XMLParsingError("word/document.xml", original_exception=e) from e
        except Exception as e:
            logger.error(f"提取 document.xml 时发生意外错误: {e}", exc_info=True)
            raise XMLParsingError("word/document.xml", original_exception=e) from e

    def read_revisions(self):
        """
        Reads inserted and deleted text from the document by parsing document.xml.

        Yields:
            dict: A dictionary containing 'type' ('inserted' or 'deleted'), 'author', 'date', and 'text'.

        Raises:
            DocumentNotLoadedError: If the document has not been loaded.
            XMLParsingError: If there's an error parsing the document XML.
        """
        if self._document is None:
            logger.error("文档未加载。请先调用 load_document()。")
            raise DocumentNotLoadedError()

        logger.info("正在从文档中读取修订。")
        document_tree = self._get_document_xml_tree()
        # _get_document_xml_tree already raises XMLParsingError if it fails

        nsmap = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'w14': 'http://schemas.microsoft.com/office/word/2010/wordml' # For some revision types
        }

        revision_count = 0
        # Find inserted text
        for ins_element in document_tree.xpath('//w:ins', namespaces=nsmap):
            author = ins_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author')
            date = ins_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date')
            text_parts = []
            for t_element in ins_element.xpath('.//w:t', namespaces=nsmap):
                if t_element.text:
                    text_parts.append(t_element.text)
            inserted_text = "".join(text_parts)
            if inserted_text:
                logger.debug(f"Read inserted text: author='{author}', date='{date}', text='{inserted_text[:50]}'")
                yield {
                    "type": "inserted",
                    "author": author,
                    "date": date,
                    "text": inserted_text
                }
                revision_count += 1

        # Find deleted text
        for del_element in document_tree.xpath('//w:del', namespaces=nsmap):
            author = del_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author')
            date = del_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date')
            text_parts = []
            for t_element in del_element.xpath('.//w:t', namespaces=nsmap):
                if t_element.text:
                    text_parts.append(t_element.text)
            deleted_text = "".join(text_parts)
            if deleted_text:
                logger.debug(f"Read deleted text: author='{author}', date='{date}', text='{deleted_text[:50]}'")
                yield {
                    "type": "deleted",
                    "author": author,
                    "date": date,
                    "text": deleted_text
                }
                revision_count += 1
        logger.info(f"Finished reading {revision_count} revisions.")

    def get_document_with_revisions_and_comments_formatted(self, include_added: bool = True,
                                                           include_deleted: bool = True,
                                                           include_comments: bool = True,
                                                           as_final_draft: bool = False,
                                                           as_original_draft: bool = False,
                                                           merge_revisions: bool = None) -> str:
        """
        Generates a formatted text representation of the document,
        including revisions (inserted/deleted text) and comments,
        using custom tags.

        Args:
            include_added: Whether to include added text with tags.
            include_deleted: Whether to include deleted text with tags.
            include_comments: Whether to include comments with tags.
            as_final_draft: If True, only includes added text (without tags) and excludes deleted text and comments.
            as_original_draft: If True, excludes added text and includes deleted text (without tags) and excludes comments.

        Returns:
            str: The formatted document text.

        Raises:
            DocumentNotLoadedError: If the document has not been loaded.
            XMLParsingError: If there's an error parsing the document or comments XML.
        """
        if self._document is None:
            logger.error("文档未加载。请先调用 load_document()。")
            raise DocumentNotLoadedError()

        logger.info("正在生成包含修订和评论的格式化文档。")
        document_tree = self._get_document_xml_tree()
        # _get_document_xml_tree already raises XMLParsingError if it fails
        comments_tree = None
        try:
            comments_tree = self._get_comments_xml_tree()
        except XMLParsingError as e:
            logger.warning(f"加载评论 XML 树时出错，将继续不带评论。异常: {e}")
            # Continue without comments if there's a parsing error for comments.xml

        nsmap = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'w14': 'http://schemas.microsoft.com/office/word/2010/wordml'
        }

        # Build a dictionary of comments for quick lookup by ID
        comments_by_id = {}
        if comments_tree and include_comments and not as_final_draft and not as_original_draft:
            for comment_element in comments_tree.xpath('//w:comment', namespaces=nsmap):
                comment_id = comment_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                comment_text_parts = []
                for p_element in comment_element.xpath('.//w:p', namespaces=nsmap):
                    for t_element in p_element.xpath('.//w:t', namespaces=nsmap):
                        if t_element.text:
                            comment_text_parts.append(t_element.text)
                comments_by_id[comment_id] = "".join(comment_text_parts)

        # Determine the actual merge_revisions setting
        _merge_revisions = self._merge_revisions if merge_revisions is None else merge_revisions

        formatted_text_lines = []

        # Iterate through paragraphs in the document XML
        for p_element in document_tree.xpath('//w:body/w:p', namespaces=nsmap):
            paragraph_content = []
            # Keep track of open comment ranges
            open_comment_ids = set()

            current_revision_type = None
            current_revision_text = ""

            # Helper to flush accumulated revision text
            def flush_revision_text():
                nonlocal current_revision_type, current_revision_text
                if current_revision_text:
                    if current_revision_type == "inserted":
                        paragraph_content.append(f"{self._added_tag_start}{current_revision_text}{self._added_tag_end}")
                    elif current_revision_type == "deleted":
                        paragraph_content.append(f"{self._deleted_tag_start}{current_revision_text}{self._deleted_tag_end}")
                    current_revision_text = ""
                    current_revision_type = None

            for child in p_element.iterchildren():
                tag = etree.QName(child.tag).localname
                text_content = ""

                # Extract text from w:t (text run) elements
                if tag == 'r': # w:r (run)
                    # Flush any pending revision text before processing normal run content
                    if _merge_revisions and current_revision_text:
                        flush_revision_text()

                    for t_element in child.xpath('.//w:t', namespaces=nsmap):
                        if t_element.text:
                            text_content += t_element.text

                    # Check for inserted/deleted runs within a normal run
                    ins_elements = child.xpath('.//w:ins', namespaces=nsmap)
                    del_elements = child.xpath('.//w:del', namespaces=nsmap)

                    if ins_elements and include_added and not as_original_draft:
                        for ins_element in ins_elements:
                            ins_text_parts = []
                            for t_elem in ins_element.xpath('.//w:t', namespaces=nsmap):
                                if t_elem.text:
                                    ins_text_parts.append(t_elem.text)
                            if ins_text_parts:
                                if _merge_revisions:
                                    if current_revision_type == "inserted":
                                        current_revision_text += ''.join(ins_text_parts)
                                    else:
                                        flush_revision_text()
                                        current_revision_type = "inserted"
                                        current_revision_text = ''.join(ins_text_parts)
                                else:
                                    text_content += f"{self._added_tag_start}{''.join(ins_text_parts)}{self._added_tag_end}"
                    elif as_final_draft or not include_added:
                        for ins_element in ins_elements:
                            ins_text_parts = []
                            for t_elem in ins_element.xpath('.//w:t', namespaces=nsmap):
                                if t_elem.text:
                                    ins_text_parts.append(t_elem.text)
                            if ins_text_parts:
                                text_content += ''.join(ins_text_parts)

                    if del_elements and include_deleted and not as_final_draft:
                        for del_element in del_elements:
                            del_text_parts = []
                            for t_elem in del_element.xpath('.//w:t', namespaces=nsmap):
                                if t_elem.text:
                                    del_text_parts.append(t_elem.text)
                            if del_text_parts:
                                if _merge_revisions:
                                    if current_revision_type == "deleted":
                                        current_revision_text += ''.join(del_text_parts)
                                    else:
                                        flush_revision_text()
                                        current_revision_type = "deleted"
                                        current_revision_text = ''.join(del_text_parts)
                                else:
                                    text_content += f"{self._deleted_tag_start}{''.join(del_text_parts)}{self._deleted_tag_end}"
                    elif as_original_draft or not include_deleted:
                        for del_element in del_elements:
                            del_text_parts = []
                            for t_elem in del_element.xpath('.//w:t', namespaces=nsmap):
                                if t_elem.text:
                                    del_text_parts.append(t_elem.text)
                            if del_text_parts:
                                text_content += ''.join(del_text_parts)

                elif tag == 'ins': # w:ins (inserted text)
                    if include_added and not as_original_draft:
                        text_parts = []
                        for t_element in child.xpath('.//w:t', namespaces=nsmap):
                            if t_element.text:
                                text_parts.append(t_element.text)
                        if text_parts:
                            if _merge_revisions:
                                if current_revision_type == "inserted":
                                    current_revision_text += ''.join(text_parts)
                                else:
                                    flush_revision_text()
                                    current_revision_type = "inserted"
                                    current_revision_text = ''.join(text_parts)
                            else:
                                text_content = f"{self._added_tag_start}{''.join(text_parts)}{self._added_tag_end}"
                    elif as_final_draft or not include_added:
                        text_parts = []
                        for t_element in child.xpath('.//w:t', namespaces=nsmap):
                            if t_element.text:
                                text_parts.append(t_element.text)
                        if text_parts:
                            text_content = ''.join(text_parts)
                elif tag == 'del': # w:del (deleted text)
                    if include_deleted and not as_final_draft:
                        text_parts = []
                        # Look for w:delText directly within w:del or within w:r inside w:del
                        for del_text_element in child.xpath('.//w:delText', namespaces=nsmap):
                            if del_text_element.text:
                                text_parts.append(del_text_element.text)
                        for t_element in child.xpath('.//w:t', namespaces=nsmap): # Also check for w:t within w:del
                            if t_element.text:
                                text_parts.append(t_element.text)
                        if text_parts:
                            if _merge_revisions:
                                if current_revision_type == "deleted":
                                    current_revision_text += ''.join(text_parts)
                                else:
                                    flush_revision_text()
                                    current_revision_type = "deleted"
                                    current_revision_text = ''.join(text_parts)
                            else:
                                text_content = f"{self._deleted_tag_start}{''.join(text_parts)}{self._deleted_tag_end}"
                    elif as_original_draft or not include_deleted:
                        text_parts = []
                        for del_text_element in child.xpath('.//w:delText', namespaces=nsmap):
                            if del_text_element.text:
                                text_parts.append(del_text_element.text)
                        for t_element in child.xpath('.//w:t', namespaces=nsmap):
                            if t_element.text:
                                text_parts.append(t_element.text)
                        if text_parts:
                            text_content = ''.join(text_parts)
                elif tag == 'commentRangeStart' and include_comments and not as_final_draft and not as_original_draft:
                    # Flush any pending revision text before a comment range
                    if _merge_revisions and current_revision_text:
                        flush_revision_text()
                    comment_id = child.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                    if comment_id in comments_by_id:
                        open_comment_ids.add(comment_id)
                elif tag == 'commentRangeEnd' and include_comments and not as_final_draft and not as_original_draft:
                    # Flush any pending revision text before a comment range
                    if _merge_revisions and current_revision_text:
                        flush_revision_text()
                    comment_id = child.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                    if comment_id in open_comment_ids:
                        comment_text = comments_by_id.get(comment_id, "Comment Not Found")
                        paragraph_content.append(f"{self._comment_tag_start}{comment_text}{self._comment_tag_end}")
                        open_comment_ids.remove(comment_id)

                if text_content: # Only append if there's actual text
                    # If not merging, or if it's normal text, append directly
                    if not _merge_revisions or (current_revision_type is None and not (tag == 'ins' or tag == 'del')):
                        paragraph_content.append(text_content)

            # After iterating through all children, flush any remaining accumulated revision text
            if _merge_revisions and current_revision_text:
                flush_revision_text()

            formatted_text_lines.append("".join(paragraph_content).strip())

        return "\n".join(formatted_text_lines)

    def get_added_text_formatted(self) -> str:
        """
        Extracts and formats only the added text from the document.

        Returns:
            str: The formatted added text.

        Raises:
            DocumentNotLoadedError: If the document has not been loaded.
            XMLParsingError: If there's an error parsing the document XML.
        """
        if self._document is None:
            logger.error("文档未加载。请先调用 load_document()。")
            raise DocumentNotLoadedError()

        logger.info("正在提取仅添加的文本。")
        document_tree = self._get_document_xml_tree()
        # _get_document_xml_tree already raises XMLParsingError if it fails

        nsmap = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'w14': 'http://schemas.microsoft.com/office/word/2010/wordml'
        }

        added_texts = []
        for ins_element in document_tree.xpath('//w:ins', namespaces=nsmap):
            text_parts = []
            for t_element in ins_element.xpath('.//w:t', namespaces=nsmap):
                if t_element.text:
                    text_parts.append(t_element.text)
            if text_parts:
                added_texts.append(f"{self._added_tag_start}{''.join(text_parts)}{self._added_tag_end}")
        return "\n".join(added_texts)

    def get_deleted_text_formatted(self) -> str:
        """
        Extracts and formats only the deleted text from the document.

        Returns:
            str: The formatted deleted text.

        Raises:
            DocumentNotLoadedError: If the document has not been loaded.
            XMLParsingError: If there's an error parsing the document XML.
        """
        if self._document is None:
            logger.error("文档未加载。请先调用 load_document()。")
            raise DocumentNotLoadedError()

        logger.info("正在提取仅删除的文本。")
        document_tree = self._get_document_xml_tree()
        # _get_document_xml_tree already raises XMLParsingError if it fails

        nsmap = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'w14': 'http://schemas.microsoft.com/office/word/2010/wordml'
        }

        deleted_texts = []
        for del_element in document_tree.xpath('//w:del', namespaces=nsmap):
            text_parts = []
            for t_element in del_element.xpath('.//w:t', namespaces=nsmap):
                if t_element.text:
                    text_parts.append(t_element.text)
            # Also check for w:delText directly within w:del
            for del_text_element in del_element.xpath('.//w:delText', namespaces=nsmap):
                if del_text_element.text:
                    text_parts.append(del_text_element.text)
            if text_parts:
                deleted_texts.append(f"{self._deleted_tag_start}{''.join(text_parts)}{self._deleted_tag_end}")
        return "\n".join(deleted_texts)

    def get_comments_formatted(self) -> str:
        """
        Extracts and formats only the comments from the document.

        Returns:
            str: The formatted comments.

        Raises:
            DocumentNotLoadedError: If the document has not been loaded.
            XMLParsingError: If there's an error parsing the comments XML.
        """
        if self._document is None:
            logger.error("文档未加载。请先调用 load_document()。")
            raise DocumentNotLoadedError()

        logger.info("正在提取仅评论。")
        comments_tree = self._get_comments_xml_tree()
        if comments_tree is None:
            return "" # _get_comments_xml_tree already logs/raises if there's a critical error

        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        comment_texts = []
        for comment_element in comments_tree.xpath('//w:comment', namespaces=nsmap):
            comment_text_parts = []
            for p_element in comment_element.xpath('.//w:p', namespaces=nsmap):
                for t_element in p_element.xpath('.//w:t', namespaces=nsmap):
                    if t_element.text:
                        comment_text_parts.append(t_element.text)
            if comment_text_parts:
                comment_texts.append(f"{self._comment_tag_start}{''.join(comment_text_parts)}{self._comment_tag_end}")
        return "\n".join(comment_texts)

    def get_final_draft(self) -> str:
        """
        Generates the final draft of the document (all additions accepted, all deletions rejected, no comments).
        """
        logger.info("Generating final draft.")
        return self.get_document_with_revisions_and_comments_formatted(as_final_draft=True, merge_revisions=False)

    def get_original_draft(self) -> str:
        """
        Generates the original draft of the document (all additions rejected, all deletions accepted, no comments).
        """
        logger.info("Generating original draft.")
        return self.get_document_with_revisions_and_comments_formatted(as_original_draft=True, merge_revisions=False)

    def replace_text(self, old_text, new_text):
        # --- ADD: Method docstring ---
        """
        Replaces the first occurrence of old_text with new_text in each paragraph.
        NOTE: This method does NOT handle revisions or comments gracefully.
        Direct manipulation on documents with complex revision/comment states
        may lead to unexpected results or data corruption.

        Args:
            old_text: The text to find.
            new_text: The text to replace with.

        Raises:
            DocumentNotLoadedError: If the document has not been loaded.
        """
        # --- END ADD ---
        if self._document is None:
            logger.error("文档未加载。请先调用 load_document()。")
            raise DocumentNotLoadedError()

        logger.info(f"尝试替换文本: '{old_text}' 为 '{new_text}'")
        replacements_made = 0
        for paragraph in self._document.paragraphs:
            # Simple string replacement on paragraph text. This is brittle
            # as runs within a paragraph can have different formatting
            # and complex runs/revisions will not be handled correctly.
            if old_text in paragraph.text:
                # A more robust replacement requires iterating through runs
                # and potentially deleting/adding runs, which is much more complex
                # and especially difficult with revisions.
                # This simple example just shows a basic idea but is NOT suitable
                # for documents with revisions.
                original_text = paragraph.text
                paragraph.text = original_text.replace(old_text, new_text, 1) # Replace first occurrence
                replacements_made += 1
                logger.debug(f"在段落中替换了文本。原始 (前50个字符): '{original_text[:50]}', 新 (前50个字符): '{paragraph.text[:50]}'")

        if replacements_made > 0:
             logger.info(f"完成尝试文本替换。进行了 {replacements_made} 次替换。")
        else:
             logger.warning(f"完成尝试文本替换。未找到 '{old_text}' 的任何出现。")

    def save_document(self, output_filepath):
        # --- ADD: Method docstring ---
        """
        Saves the modified document to a new filepath.

        Args:
            output_filepath: The path to save the modified .docx file.

        Raises:
            DocumentNotLoadedError: If the document has not been loaded.
            DocumentSavingError: If there's an error saving the document.
        """
        # --- END ADD ---
        if self._document is None:
            logger.error("文档未加载。请先调用 load_document()。")
            raise DocumentNotLoadedError()

        try:
            self._document.save(output_filepath)
            logger.info(f"成功将文档保存到: {output_filepath}")
        except Exception as e:
            logger.error(f"保存文档到 {output_filepath} 错误。异常: {e}", exc_info=True)
            raise DocumentSavingError(output_filepath, original_exception=e) from e
