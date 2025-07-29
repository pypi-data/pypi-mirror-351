import os
import argparse
import git
import shutil
from pathlib import Path
from docx import Document
import tempfile
import logging
import stat
from typing import Generator, Optional, List, Union

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class UnifiedSourceToTextConverter:
    def __init__(self, input_path: str, output_path: str, output_type: str = "txt", verbose: bool = False, exclude_hidden: bool = False, include_exts: Optional[List[str]] = None, branch: Optional[str] = None) -> None:
        self.input_path: str = input_path
        self.output_path: str = output_path
        self.output_type: str = output_type.lower()
        self.verbose: bool = verbose
        self.exclude_hidden: bool = exclude_hidden
        self.include_exts: Optional[List[str]] = [ext.lower() for ext in include_exts] if include_exts else None
        self.branch: Optional[str] = branch
        self.temp_folder_path: Optional[str] = None
        if self.verbose:
            logger.setLevel(logging.DEBUG)
        else:
            pass

    def _is_hidden_file(self, file_path: Union[str, Path]) -> bool:
        path = Path(file_path)
        for part in path.parts:
            if part.startswith('.') or part.startswith('__'):
                return True
        return False

    def _is_included_file(self, file_path: str) -> bool:
        if self.include_exts is None:
            return True
        ext = os.path.splitext(file_path)[1].lower()
        return ext in self.include_exts

    def _parse_folder(self, folder_path: str) -> str:
        tree: str = ""
        for root, dirs, files in os.walk(folder_path):
            if self.exclude_hidden:
                dirs[:] = [d for d in dirs if not self._is_hidden_file(os.path.join(root, d))]

            level = root.replace(folder_path, '').count(os.sep)
            indent = ' ' * 4 * (level)
            if self.exclude_hidden and self._is_hidden_file(os.path.basename(root)):
                continue
            tree += '{}{}/\n'.format(indent, os.path.basename(root))
            subindent = ' ' * 4 * (level + 1)
            if self.exclude_hidden:
                files = [f for f in files if not self._is_hidden_file(os.path.join(root, f))]
            if self.include_exts is not None:
                files = [f for f in files if self._is_included_file(f)]
            for f in files:
                tree += '{}{}\n'.format(subindent, f)

        if self.verbose:
            logger.debug(f"The file tree to be processed:\n{tree}")

        return tree

    def _get_file_contents(self, file_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            logger.warning(f"Could not decode file {file_path} with utf-8. Skipping content.")
            return ""
        except Exception as e:
            logger.warning(f"Error reading file {file_path}: {e}")
            return ""

    def _process_files(self, path: str) -> Generator[str, None, None]:
        for root, _, files in os.walk(path):
            if '/.git' in root:
                if self.verbose:
                    logger.debug(f"Ignoring .git directory: {root}")
                continue

            for file in files:
                file_path = os.path.join(root, file)
                if '/.git/' in file_path:
                    if self.verbose:
                        logger.debug(f"Ignoring file in .git directory: {file_path}")
                    continue

                if self.exclude_hidden and self._is_hidden_file(os.path.abspath(file_path)):
                    if self.verbose:
                        logger.debug(f"Ignoring hidden file {file_path}")
                    continue
                if self.include_exts is not None and not self._is_included_file(file):
                    if self.verbose:
                        logger.debug(f"Skipping file due to extension filter: {file_path}")
                    continue
                try:
                    if self.verbose:
                        logger.debug(f"Processing: {file_path}")
                    file_content = self._get_file_contents(file_path)
                    yield f"\n\n{file_path}\n"
                    yield f"File type: {os.path.splitext(file_path)[1]}\n"
                    yield file_content
                    yield f"\n\n{'-' * 50}\nFile End\n{'-' * 50}\n"
                except Exception as e:
                    logger.error(f"Couldn't process {file_path}: {e}")

    def get_text(self) -> dict:
        """
        Returns structured data containing folder structure and file contents.
        For 'txt' output, returns a dict with 'folder_structure' and 'files' list.
        For 'docx' output, returns the same structured data for easier processing.
        """
        folder_structure: str = ""
        files_data: list = []
        if self.is_github_repo():
            self._clone_github_repo()
            folder_structure = self._parse_folder(self.temp_folder_path)
            for root, _, files in os.walk(self.temp_folder_path):
                for file in files:
                    file_path = os.path.join(root, file)
                    if self.exclude_hidden and self._is_hidden_file(file_path):
                        continue
                    if self.include_exts is not None and not self._is_included_file(file):
                        continue
                    content = self._get_file_contents(file_path)
                    relative_path = os.path.relpath(file_path, self.temp_folder_path)
                    files_data.append({"path": relative_path, "content": content})
            # Do not clean up here to allow reuse in get_file for docx
        else:
            folder_structure = self._parse_folder(self.input_path)
            for root, _, files in os.walk(self.input_path):
                for file in files:
                    file_path = os.path.join(root, file)
                    if self.exclude_hidden and self._is_hidden_file(file_path):
                        continue
                    if self.include_exts is not None and not self._is_included_file(file):
                        continue
                    content = self._get_file_contents(file_path)
                    relative_path = os.path.relpath(file_path, self.input_path)
                    files_data.append({"path": relative_path, "content": content})

        return {"folder_structure": folder_structure, "files": files_data}

    def get_file(self) -> None:
        if self.is_github_repo() and self.temp_folder_path is None:
            self._clone_github_repo()
        if self.output_type == "txt":
            with open(self.output_path, "w", encoding='utf-8') as file:
                folder_structure_header = "Folder Structure"
                file_contents_header = "File Contents"
                delimiter = "-" * 50
                file.write(f"{folder_structure_header}\n{delimiter}\n")
                folder_structure = self._parse_folder(self.input_path if not self.is_github_repo() else self.temp_folder_path)
                file.write(f"{folder_structure}\n\n")
                file.write(f"{file_contents_header}\n{delimiter}\n")
                path_to_process = self.input_path if not self.is_github_repo() else self.temp_folder_path
                for content in self._process_files(path_to_process):
                    file.write(content)
            if self.is_github_repo():
                self.clean_up_temp_folder()
        elif self.output_type == "docx":
            data = self.get_text()
            doc = Document()
            doc.add_heading("Folder Structure", level=1)
            doc.add_paragraph(data["folder_structure"])
            for file_info in data["files"]:
                doc.add_heading(file_info["path"], level=2)
                doc.add_paragraph(file_info["content"])
            doc.save(self.output_path)
            if self.is_github_repo():
                self.clean_up_temp_folder()
        else:
            raise ValueError("Invalid output type. Supported types: txt, docx")

    def _clone_github_repo(self) -> None:
        try:
            self.temp_folder_path = tempfile.mkdtemp(prefix="github_repo_")
            if self.verbose:
                logger.info(f"Cloning GitHub repository {self.input_path} into {self.temp_folder_path}")
            if self.branch:
                git.Repo.clone_from(self.input_path, self.temp_folder_path, branch=self.branch)
            else:
                git.Repo.clone_from(self.input_path, self.temp_folder_path)
            if self.verbose:
                logger.info("GitHub repository cloned successfully.")
        except git.exc.GitCommandError as e:
            logger.error(f"Git clone command failed: {e}")
            self.clean_up_temp_folder()
            raise
        except OSError as e:
            logger.error(f"OS error during cloning: {e}")
            self.clean_up_temp_folder()
            raise

    def is_github_repo(self) -> bool:
        # Modified to detect GitHub URLs with tokens as well
        return (
            self.input_path.startswith("https://github.com/")
            or self.input_path.startswith("git@github.com:")
            or "github.com" in self.input_path
        )

    def _handle_remove_readonly(self, func, path, exc) -> None:
        os.chmod(path, stat.S_IWRITE)
        logger.debug(f"Attempting to remove read-only file: {path}")
        try:
            func(path)
        except Exception as e:
            logger.error(f"Failed to remove read-only file {path}: {e}")

    def clean_up_temp_folder(self) -> None:
        if self.temp_folder_path:
            if self.verbose:
                logger.info(f"Cleaning up temporary folder {self.temp_folder_path}")
            try:
                shutil.rmtree(self.temp_folder_path, onerror=self._handle_remove_readonly)
            except Exception as e:
                logger.warning(f"Failed to clean up temporary folder {self.temp_folder_path}: {e}")
            self.temp_folder_path = None

    def main() -> None:
        parser = argparse.ArgumentParser(description="Generate text from codebase.")
        parser.add_argument("--input", help="Input path (folder or GitHub URL)", required=True)
        parser.add_argument("--output", help="Output file path", required=True)
        parser.add_argument("--output-type", help="Output file type (txt or docx)", default="txt")
        parser.add_argument("--exclude-hidden", help="Exclude hidden files and folders", action="store_true")
        parser.add_argument("--include-ext", help="Comma-separated list of file extensions to include (e.g. .py,.js)")
        parser.add_argument("--verbose", help="Show useful information", action="store_true")
        parser.add_argument("--branch", help="GitHub branch to clone", default=None)
        args = parser.parse_args()

        include_exts = None
        if args.include_ext:
            include_exts = [ext.strip().lower() for ext in args.include_ext.split(",")]

        converter = UnifiedSourceToTextConverter(
            input_path=args.input,
            output_path=args.output,
            output_type=args.output_type,
            verbose=args.verbose,
            exclude_hidden=args.exclude_hidden,
            include_exts=include_exts,
            branch=args.branch
        )
        converter.get_file()

if __name__ == "__main__":
    main()
