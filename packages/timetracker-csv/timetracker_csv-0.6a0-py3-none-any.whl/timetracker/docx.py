"""Generate a Microsoft Word document containing a table of data"""

__copyright__ = 'Copyright (C) 2025-present, DV Klopfenstein, PhD. All rights reserved.'
__author__ = "DV Klopfenstein, PhD"

##from os import remove
#from os.path import exists
##from os.path import basename
##from os.path import join
##from os.path import abspath
##from os.path import dirname
##from os.path import normpath
##https://python-docx.readthedocs.io/en/latest/
from docx import Document
from docx.shared import Inches
#from docx.shared import Inches
#from datetime import timedelta
#from datetime import datetime
#from logging import debug
#
#from timetracker.utils import orange
#from timetracker.consts import DIRTRK
##from timetracker.cfg.utils import get_username


class WordDoc:
    """Generate a Microsoft Word document containing a table of data"""
    # pylint: disable=too-few-public-methods



    def __init__(self, time_formatted):
        self.nttext = time_formatted

    def write_doc(self, fout_docx):
        """Write a report into a Microsoft Word document"""
        document = Document()

        document.add_heading('Document Title', 0)

        #p = document.add_paragraph('A plain paragraph having some ')
        #p.add_run('bold').bold = True
        #p.add_run(' and some ')
        #p.add_run('italic.').italic = True

        #document.add_heading('Heading, level 1', level=1)
        #document.add_paragraph('Intense quote', style='Intense Quote')

        #document.add_paragraph(
        #    'first item in unordered list', style='List Bullet'
        #)
        #document.add_paragraph(
        #    'first item in ordered list', style='List Number'
        #)

        #document.add_picture('monty-truth.png', width=Inches(1.25))

        self._add_table(document)
        document.add_page_break()

        document.save(fout_docx)
        print(f'  WROTE: {fout_docx}')

    def _get_headers(self):
        """Get the number of rows in the timetracking data (self.nttext must have data)"""
        return self.nttext[0]._fields

    def _get_nrows(self):
        """Get the number of rows in the timetracking data (self.nttext must have data)"""
        return len(self.nttext)

    def _get_ncols(self):
        """Get the number of rows in the timetracking data (self.nttext must have data)"""
        return len(self.nttext[0])

    def _add_table(self, doc):
        """Add a table containing timetracking data to a Word document"""
        if not self.nttext:
            return
        wdct = {
            'Day': Inches(.20),
            'Date': Inches(1),
            'Duration': Inches(.5),
            'Total': Inches(.5),
            'Price': Inches(.5),
            'Description': Inches(3),
        }
        table = doc.add_table(rows=1, cols=self._get_ncols(), style='Table Grid')
        hdrs = self._get_headers()
        for hdr, cell in zip(hdrs, table.rows[0].cells):
            if hdr in wdct:
                cell.width = wdct[hdr]
            cell.text = hdr
        for ntd in self.nttext:
            row_cells = table.add_row().cells
            for hdr, cell, val in zip(hdrs, row_cells, list(ntd)):
                if hdr in wdct:
                    cell.width = wdct[hdr]
                cell.text = val


# Copyright (C) 2025-present, DV Klopfenstein, PhD. All rights reserved.
