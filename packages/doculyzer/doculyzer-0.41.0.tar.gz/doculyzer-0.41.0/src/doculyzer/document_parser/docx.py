"""
DOCX document parser module for the document pointer system.

This module parses DOCX documents into structured elements with comprehensive date extraction.
"""

import json
import logging
import os
from typing import Dict, Any, Optional, List, Union

from ..relationships import RelationshipType
from ..storage import ElementType

try:
    import docx
    from docx.document import Document as DocxDocument
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    DOCX_AVAILABLE = True
except ImportError:
    docx = None
    DocxDocument = None
    CT_Tbl = None
    CT_P = None
    Table = None
    Paragraph = None
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available. Install with 'pip install python-docx' to use DOCX parser")

from bs4 import BeautifulSoup

from .base import DocumentParser
from .extract_dates import DateExtractor

logger = logging.getLogger(__name__)


class DocxParser(DocumentParser):
    """Parser for DOCX documents with enhanced date extraction."""

    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """Initialize the DOCX parser."""
        super().__init__(config)

        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX parsing")

        # Configuration options
        self.config = config or {}
        self.extract_comments = self.config.get("extract_comments", True)
        self.extract_headers_footers = self.config.get("extract_headers_footers", True)
        self.extract_styles = self.config.get("extract_styles", True)
        self.track_changes = self.config.get("track_changes", False)
        self.max_image_size = self.config.get("max_image_size", 1024 * 1024)  # 1MB default
        self.temp_dir = self.config.get("temp_dir", os.path.join(os.path.dirname(__file__), 'temp'))
        self.max_content_preview = self.config.get("max_content_preview", 100)

        # Date extraction configuration
        self.extract_dates = self.config.get("extract_dates", True)
        self.date_context_chars = self.config.get("date_context_chars", 50)  # Small context window
        self.min_year = self.config.get("min_year", 1900)
        self.max_year = self.config.get("max_year", 2100)
        self.fiscal_year_start_month = self.config.get("fiscal_year_start_month", 10)
        self.default_locale = self.config.get("default_locale", "US")

        # Initialize date extractor if enabled
        self.date_extractor = None
        if self.extract_dates:
            try:
                self.date_extractor = DateExtractor(
                    context_chars=self.date_context_chars,
                    min_year=self.min_year,
                    max_year=self.max_year,
                    fiscal_year_start_month=self.fiscal_year_start_month,
                    default_locale=self.default_locale
                )
                logger.debug("Date extraction enabled with comprehensive temporal analysis")
            except ImportError as e:
                logger.warning(f"Date extraction disabled: {e}")
                self.extract_dates = False

    def _resolve_element_text(self, location_data: Dict[str, Any], source_content: Optional[Union[str, bytes]] = None) -> str:
        """
        Resolve the plain text representation of a DOCX element.

        Args:
            location_data: Content location data
            source_content: Optional preloaded source content

        Returns:
            Plain text representation of the element
        """
        # DOCX content is already text-based, so we can leverage the existing _resolve_element_content method
        # We just need to clean up the returned content for certain element types

        # First, get the content using the existing method
        content = self._resolve_element_content(location_data, source_content)

        # For DOCX, the element content is already in text form without any markup
        element_type = location_data.get("type", "")

        # For most element types, the content is already in the desired plain text format
        # But we might need special handling for some types
        if element_type == ElementType.TABLE.value or element_type == ElementType.TABLE_ROW.value:
            # For tables, content from _resolve_element_content uses | as separators
            # We'll keep this format as it's already a good text representation
            return content

        elif element_type == ElementType.TABLE_CELL.value or element_type == ElementType.TABLE_HEADER.value:
            # For cells, just return the text content
            return content.strip()

        elif element_type == ElementType.HEADER.value or element_type == ElementType.PARAGRAPH.value or element_type == ElementType.LIST_ITEM.value:
            # For text elements, just return the content
            return content.strip()

        elif element_type == ElementType.COMMENT.value:
            # For comments, extract just the text without metadata
            return content.strip()

        # For other element types, just return the content as is
        return content.strip()

    def parse(self, doc_content: Dict[str, Any]) -> Dict[str, Any]:
        """
        Parse a DOCX document into structured elements with comprehensive date extraction.

        Args:
            doc_content: Document content and metadata

        Returns:
            Dictionary with document metadata, elements, relationships, extracted links, and dates
        """
        # Extract metadata from doc_content
        source_id = doc_content["id"]
        metadata = doc_content.get("metadata", {}).copy()  # Make a copy to avoid modifying original

        # Get binary content
        content = doc_content.get("content")
        if not content and "binary_path" in doc_content:
            # If we have a path but no content, load the content
            try:
                with open(doc_content["binary_path"], 'rb') as f:
                    content = f.read()
            except Exception as e:
                logger.error(f"Error loading DOCX from path: {str(e)}")
                raise

        # If we still don't have content, raise an error
        if not content:
            raise ValueError("No DOCX content provided")

        # Save content to a temporary file if needed
        binary_path = None
        try:
            # Create temp directory if it doesn't exist
            if not os.path.exists(self.temp_dir):
                os.makedirs(self.temp_dir, exist_ok=True)

            # Write content to temp file
            import uuid
            binary_path = os.path.join(self.temp_dir, f"temp_{uuid.uuid4().hex}.docx")
            with open(binary_path, 'wb') as f:
                if isinstance(content, str):
                    f.write(content.encode('utf-8'))
                else:
                    f.write(content)

            # Generate document ID if not present
            doc_id = metadata.get("doc_id", self._generate_id("doc_"))

            # Load DOCX document
            try:
                doc = docx.Document(binary_path)
            except Exception as e:
                logger.error(f"Error loading DOCX document: {str(e)}")
                raise

            # Create document record with metadata
            document = {
                "doc_id": doc_id,
                "doc_type": "docx",
                "source": source_id,
                "metadata": self._extract_document_metadata(doc, metadata),
                "content_hash": doc_content.get("content_hash", "")
            }

            # Create root element
            elements = [self._create_root_element(doc_id, source_id)]
            root_id = elements[0]["element_id"]

            # Initialize relationships list
            relationships = []

            # Parse document elements and create relationships
            elements_from_doc, relationships_from_doc = self._parse_document_elements(doc, doc_id, root_id, source_id)
            elements.extend(elements_from_doc)
            relationships.extend(relationships_from_doc)

            # Extract links from the document
            links = self._extract_links(doc, elements)

            # Extract dates from document with comprehensive temporal analysis
            element_dates = {}
            if self.extract_dates and self.date_extractor:
                try:
                    # Extract text content from the entire document for date extraction
                    full_text = self._extract_full_text(doc)

                    # Extract dates from the full document
                    if full_text.strip():
                        document_dates = self.date_extractor.extract_dates_as_dicts(full_text)
                        if document_dates:
                            element_dates[root_id] = document_dates
                            logger.debug(f"Extracted {len(document_dates)} dates from DOCX document")

                    # Extract dates from individual elements
                    for element in elements:
                        element_id = element["element_id"]
                        element_type = element["element_type"]

                        # Only extract dates from text-containing elements
                        if element_type in [ElementType.PARAGRAPH.value, "header", ElementType.LIST_ITEM.value,
                                          ElementType.TABLE_CELL.value, "table_header", ElementType.COMMENT.value]:
                            # Get the text content for this element
                            element_text = self._get_element_text_for_dates(element, doc)

                            if element_text and element_text.strip():
                                element_specific_dates = self.date_extractor.extract_dates_as_dicts(element_text)
                                if element_specific_dates:
                                    element_dates[element_id] = element_specific_dates
                                    logger.debug(f"Extracted {len(element_specific_dates)} dates from {element_type} element")

                except Exception as e:
                    logger.warning(f"Error during date extraction: {e}")

            # Add date statistics to document metadata
            if element_dates:
                total_dates = sum(len(dates) for dates in element_dates.values())
                document["metadata"]["date_extraction"] = {
                    "total_dates_found": total_dates,
                    "elements_with_dates": len(element_dates),
                    "extraction_enabled": True
                }
            else:
                document["metadata"]["date_extraction"] = {
                    "total_dates_found": 0,
                    "elements_with_dates": 0,
                    "extraction_enabled": self.extract_dates
                }

            # Return the parsed document with comprehensive date information
            result = {
                "document": document,
                "elements": elements,
                "links": links,
                "relationships": relationships
            }

            # Add dates if any were extracted
            if element_dates:
                result["element_dates"] = element_dates

            return result
        finally:
            # Clean up temporary file
            if binary_path and os.path.exists(binary_path):
                try:
                    os.remove(binary_path)
                except Exception as e:
                    logger.warning(f"Failed to delete temporary file {binary_path}: {str(e)}")

    def _extract_full_text(self, doc: DocxDocument) -> str:
        """
        Extract all text content from the document for date extraction.

        Args:
            doc: The DOCX document

        Returns:
            Full text content of the document
        """
        text_parts = []

        # Extract text from all paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = " ".join(p.text for p in cell.paragraphs).strip()
                    if cell_text:
                        text_parts.append(cell_text)

        # Extract text from headers and footers if enabled
        if self.extract_headers_footers:
            try:
                for section in doc.sections:
                    # Process headers
                    for header_type in ['first_page_header', 'header', 'even_page_header']:
                        header = getattr(section, header_type)
                        if header and header.is_linked_to_previous is False:
                            header_text = ""
                            for paragraph in header.paragraphs:
                                header_text += paragraph.text + "\n"
                            if header_text.strip():
                                text_parts.append(header_text.strip())

                    # Process footers
                    for footer_type in ['first_page_footer', 'footer', 'even_page_footer']:
                        footer = getattr(section, footer_type)
                        if footer and footer.is_linked_to_previous is False:
                            footer_text = ""
                            for paragraph in footer.paragraphs:
                                footer_text += paragraph.text + "\n"
                            if footer_text.strip():
                                text_parts.append(footer_text.strip())
            except Exception as e:
                logger.warning(f"Error extracting text from headers/footers: {e}")

        # Extract text from comments if enabled
        if self.extract_comments:
            try:
                # Get comments part if it exists
                if doc.part.package.parts:
                    for rel_type, parts in doc.part.package.rels.items():
                        if 'comments' in rel_type.lower():
                            for rel_id, rel in parts.items():
                                if hasattr(rel, 'target_part') and rel.target_part:
                                    comments_xml = rel.target_part.blob
                                    if comments_xml:
                                        soup = BeautifulSoup(comments_xml, 'xml')
                                        for comment in soup.find_all('comment'):
                                            comment_text = comment.get_text().strip()
                                            if comment_text:
                                                text_parts.append(comment_text)
            except Exception as e:
                logger.warning(f"Error extracting text from comments: {e}")

        return "\n".join(text_parts)

    @staticmethod
    def _get_element_text_for_dates(element: Dict[str, Any], doc: DocxDocument) -> str:
        """
        Get the text content of a specific element for date extraction.

        Args:
            element: Element dictionary
            doc: The DOCX document

        Returns:
            Text content of the element
        """
        try:
            # Parse the content location to get element details
            content_location = json.loads(element["content_location"])
            element_type = content_location.get("type", "")

            if element_type == ElementType.PARAGRAPH.value:
                index = content_location.get("index", 0)
                if 0 <= index < len(doc.paragraphs):
                    return doc.paragraphs[index].text.strip()

            elif element_type == "header":
                # For headers, we already have the text in the content_preview
                return element.get("content_preview", "")

            elif element_type == ElementType.LIST_ITEM.value:
                index = content_location.get("index", 0)
                if 0 <= index < len(doc.paragraphs):
                    return doc.paragraphs[index].text.strip()

            elif element_type in [ElementType.TABLE_CELL.value, "table_header"]:
                table_index = content_location.get("table_index", 0)
                row = content_location.get("row", 0)
                col = content_location.get("col", 0)

                tables = [t for t in doc._body._body.iterchildren() if isinstance(t, CT_Tbl)]
                if 0 <= table_index < len(tables):
                    table = Table(tables[table_index], doc._body)
                    if 0 <= row < len(table.rows) and 0 <= col < len(table.rows[row].cells):
                        cell = table.rows[row].cells[col]
                        return " ".join(p.text for p in cell.paragraphs).strip()

            elif element_type == ElementType.COMMENT.value:
                # For comments, we can use the text from metadata
                return element.get("metadata", {}).get("text", "")

            elif element_type in [ElementType.PAGE_HEADER.value, ElementType.PAGE_FOOTER.value]:
                # For headers/footers, we can use the text from metadata
                return element.get("metadata", {}).get("text", "")

        except Exception as e:
            logger.warning(f"Error getting element text for dates: {e}")

        # Fallback to content preview
        return element.get("content_preview", "")

    def _resolve_element_content(self, location_data: Dict[str, Any],
                                 source_content: Optional[Union[str, bytes]]) -> str:
        """
        Resolve content for specific DOCX element types.

        Args:
            location_data: Content location data
            source_content: Optional pre-loaded source content

        Returns:
            Resolved content string
        """
        source = location_data.get("source", "")
        element_type = location_data.get("type", "")

        # Load the document if source content is not provided
        doc = None
        temp_file = None
        try:
            if source_content is None:
                # Check if source is a file path
                if os.path.exists(source):
                    try:
                        doc = docx.Document(source)
                    except Exception as e:
                        raise ValueError(f"Error loading DOCX document: {str(e)}")
                else:
                    raise ValueError(f"Source file not found: {source}")
            else:
                # Save content to a temporary file
                if not os.path.exists(self.temp_dir):
                    os.makedirs(self.temp_dir, exist_ok=True)

                import uuid
                temp_file = os.path.join(self.temp_dir, f"temp_{uuid.uuid4().hex}.docx")
                with open(temp_file, 'wb') as f:
                    if isinstance(source_content, str):
                        f.write(source_content.encode('utf-8'))
                    else:
                        f.write(source_content)

                # Load the document
                try:
                    doc = docx.Document(temp_file)
                except Exception as e:
                    raise ValueError(f"Error loading DOCX document: {str(e)}")

            # Handle different element types
            if element_type == ElementType.PARAGRAPH.value:
                # Extract paragraph by index
                index = location_data.get("index", 0)
                if 0 <= index < len(doc.paragraphs):
                    return doc.paragraphs[index].text
                return ""

            elif element_type == ElementType.HEADER.value:
                # Extract header by level and/or text
                level = location_data.get("level")
                text = location_data.get("text", "")

                # Find headers with appropriate level
                headers = []
                for para in doc.paragraphs:
                    style = para.style.name.lower() if para.style else ""
                    para_level = None

                    if style == 'title':
                        para_level = 1
                    elif style == 'subtitle':
                        para_level = 2
                    elif style.startswith('heading '):
                        try:
                            para_level = int(style.split(' ')[1])
                        except (IndexError, ValueError):
                            pass

                    if (level is None or para_level == level) and (not text or text in para.text):
                        headers.append(para)

                # Return the first matching header
                if headers:
                    return headers[0].text
                return ""

            elif element_type == ElementType.TABLE.value:
                # Extract table by index
                table_index = location_data.get("index", 0)
                tables = [t for t in doc._body._body.iterchildren() if isinstance(t, CT_Tbl)]

                if 0 <= table_index < len(tables):
                    table = Table(tables[table_index], doc._body)
                    result = []
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = " ".join(p.text for p in cell.paragraphs).strip()
                            row_text.append(cell_text)
                        result.append(" | ".join(row_text))
                    return "\n".join(result)
                return ""

            elif element_type == ElementType.TABLE_ROW.value:
                # Extract table row
                table_index = location_data.get("table_index", 0)
                row = location_data.get("row", 0)

                tables = [t for t in doc._body._body.iterchildren() if isinstance(t, CT_Tbl)]
                if 0 <= table_index < len(tables):
                    table = Table(tables[table_index], doc._body)
                    if 0 <= row < len(table.rows):
                        row_text = []
                        for cell in table.rows[row].cells:
                            cell_text = " ".join(p.text for p in cell.paragraphs).strip()
                            row_text.append(cell_text)
                        return " | ".join(row_text)
                return ""

            elif element_type == ElementType.TABLE_CELL.value or element_type == ElementType.TABLE_HEADER.value:
                # Extract table cell
                table_index = location_data.get("table_index", 0)
                row = location_data.get("row", 0)
                col = location_data.get("col", 0)

                tables = [t for t in doc._body._body.iterchildren() if isinstance(t, CT_Tbl)]
                if 0 <= table_index < len(tables):
                    table = Table(tables[table_index], doc._body)
                    if 0 <= row < len(table.rows) and 0 <= col < len(table.rows[row].cells):
                        cell = table.rows[row].cells[col]
                        return " ".join(p.text for p in cell.paragraphs).strip()
                return ""

            elif element_type == ElementType.PAGE_HEADER.value or element_type == ElementType.PAGE_FOOTER.value:
                # Extract header/footer
                section = location_data.get("section", 0)
                header_type = location_data.get("header_type", "")
                footer_type = location_data.get("footer_type", "")

                if 0 <= section < len(doc.sections):
                    section_obj = doc.sections[section]

                    if element_type == ElementType.PAGE_HEADER.value and header_type:
                        attr_name = header_type.replace(' ', '_')
                        header = getattr(section_obj, attr_name, None)
                        if header and header.is_linked_to_previous is False:
                            return "\n".join(p.text for p in header.paragraphs)

                    elif element_type == ElementType.PAGE_FOOTER.value and footer_type:
                        attr_name = footer_type.replace(' ', '_')
                        footer = getattr(section_obj, attr_name, None)
                        if footer and footer.is_linked_to_previous is False:
                            return "\n".join(p.text for p in footer.paragraphs)
                return ""

            elif element_type == ElementType.COMMENT.value:
                # Extract comment by ID
                comment_id = location_data.get("comment_id", "")

                # Python-docx doesn't have direct API for comments
                # This is a simplified approach that may not work for all documents
                if doc.part.package.parts:
                    for rel_type, parts in doc.part.package.rels.items():
                        if 'comments' in rel_type.lower():
                            for rel_id, rel in parts.items():
                                if hasattr(rel, 'target_part') and rel.target_part:
                                    comments_xml = rel.target_part.blob
                                    if comments_xml:
                                        soup = BeautifulSoup(comments_xml, 'xml')
                                        for comment in soup.find_all('comment'):
                                            if comment.get('id', '') == comment_id:
                                                return comment.get_text().strip()
                return ""

            elif element_type == ElementType.BODY.value:
                # Return all paragraphs in the document body
                return "\n".join(p.text for p in doc.paragraphs)

            else:
                # For other element types or if no specific handler,
                # return full document text
                return "\n".join(p.text for p in doc.paragraphs)

        finally:
            # Clean up temporary file
            if temp_file and os.path.exists(temp_file):
                try:
                    os.remove(temp_file)
                except Exception as e:
                    logger.warning(f"Failed to delete temporary file {temp_file}: {str(e)}")

    def supports_location(self, content_location: Dict[str, Any]) -> bool:
        """
        Check if this parser supports resolving the given location.

        Args:
            content_location: Content location pointer

        Returns:
            True if supported, False otherwise
        """
        try:
            location_data = content_location
            source = location_data.get("source", "")

            # Check if source exists and is a file
            if not os.path.exists(source) or not os.path.isfile(source):
                return False

            # Check file extension for DOCX
            _, ext = os.path.splitext(source.lower())
            return ext in ['.docx']

        except (json.JSONDecodeError, TypeError):
            return False

    def _extract_document_metadata(self, doc: DocxDocument, base_metadata: Dict[str, Any]) -> Dict[str, Any]:
        """
        Extract metadata from DOCX document.

        Args:
            doc: The DOCX document
            base_metadata: Base metadata from content source

        Returns:
            Dictionary of document metadata
        """
        # Combine base metadata with document core properties
        metadata = base_metadata.copy()

        try:
            # Get core properties
            core_props = doc.core_properties

            # Add core properties to metadata
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.created:
                metadata["created"] = core_props.created.timestamp() if hasattr(core_props.created,
                                                                                'timestamp') else str(
                    core_props.created)
            if core_props.modified:
                metadata["modified"] = core_props.modified.timestamp() if hasattr(core_props.modified,
                                                                                  'timestamp') else str(
                    core_props.modified)
            if core_props.last_modified_by:
                metadata["last_modified_by"] = core_props.last_modified_by
            if core_props.keywords:
                metadata["keywords"] = core_props.keywords
            if core_props.subject:
                metadata["subject"] = core_props.subject
            if core_props.comments:
                metadata["comments"] = core_props.comments
            if core_props.category:
                metadata["category"] = core_props.category

            # Add document statistics
            metadata["page_count"] = self._estimate_page_count(doc)
            metadata["paragraph_count"] = len(doc.paragraphs)
            metadata["word_count"] = self._count_words(doc)

            # Add style information if enabled
            if self.extract_styles:
                metadata["styles"] = self._extract_styles_info(doc)

        except Exception as e:
            logger.warning(f"Error extracting document metadata: {str(e)}")

        return metadata

    def _parse_document_elements(self, doc: DocxDocument, doc_id: str, parent_id: str, source_id: str) -> tuple[
        List[Dict[str, Any]], List[Dict[str, Any]]]:
        """
        Parse DOCX document into structured elements and create relationships.

        Args:
            doc: The DOCX document
            doc_id: Document ID
            parent_id: Parent element ID
            source_id: Source identifier

        Returns:
            Tuple of (list of elements, list of relationships)
        """
        elements = []
        relationships = []
        section_stack = [{"id": parent_id, "level": 0}]

        # Extract headers and footers if enabled
        if self.extract_headers_footers:
            header_elements, header_relationships = self._extract_headers_footers(doc, doc_id, parent_id, source_id)
            elements.extend(header_elements)
            relationships.extend(header_relationships)

        # Process document body
        body_id = self._generate_id("body_")
        body_element = {
            "element_id": body_id,
            "doc_id": doc_id,
            "element_type": ElementType.BODY.value,
            "parent_id": parent_id,
            "content_preview": "Document body",
            "content_location": json.dumps({
                "source": source_id,
                "type": ElementType.BODY.value
            }),
            "content_hash": "",
            "metadata": {}
        }
        elements.append(body_element)

        # Create relationship from root to body
        contains_relationship = {
            "relationship_id": self._generate_id("rel_"),
            "source_id": parent_id,
            "target_id": body_id,
            "relationship_type": RelationshipType.CONTAINS.value,
            "metadata": {
                "confidence": 1.0
            }
        }
        relationships.append(contains_relationship)

        # Create inverse relationship
        contained_by_relationship = {
            "relationship_id": self._generate_id("rel_"),
            "source_id": body_id,
            "target_id": parent_id,
            "relationship_type": RelationshipType.CONTAINED_BY.value,
            "metadata": {
                "confidence": 1.0
            }
        }
        relationships.append(contained_by_relationship)

        current_parent = body_id

        # Process all block-level elements in the document
        for i, block in enumerate(self._iter_block_items(doc)):
            if isinstance(block, Paragraph):
                # Process paragraph
                para_element = self._process_paragraph(block, i, doc_id, current_parent, source_id)

                # Skip empty paragraphs
                if not para_element:
                    continue

                # Check for headings
                style = block.style.name.lower() if block.style else ""
                if style.startswith('heading ') or style == 'title' or style == 'subtitle':
                    # This is a heading paragraph
                    level = 1  # Default level

                    if style == 'title':
                        level = 1
                    elif style == 'subtitle':
                        level = 2
                    elif style.startswith('heading '):
                        try:
                            level = int(style.split(' ')[1])
                        except (IndexError, ValueError):
                            pass

                    # Update element type and metadata
                    para_element["element_type"] = "header"
                    para_element["metadata"]["level"] = level

                    # Update section stack and current parent
                    while section_stack[-1]["level"] >= level:
                        section_stack.pop()

                    current_parent = section_stack[-1]["id"]
                    para_element["parent_id"] = current_parent

                    # Add to section stack
                    section_stack.append({"id": para_element["element_id"], "level": level})
                    current_parent = para_element["element_id"]

                elements.append(para_element)

                # Create relationship from parent to paragraph/header
                element_id = para_element["element_id"]
                contains_relationship = {
                    "relationship_id": self._generate_id("rel_"),
                    "source_id": para_element["parent_id"],
                    "target_id": element_id,
                    "relationship_type": RelationshipType.CONTAINS.value if para_element[
                                                                                "element_type"] == "header" else RelationshipType.CONTAINS_TEXT.value,
                    "metadata": {
                        "confidence": 1.0,
                        "index": i
                    }
                }
                relationships.append(contains_relationship)

                # Create inverse relationship
                contained_by_relationship = {
                    "relationship_id": self._generate_id("rel_"),
                    "source_id": element_id,
                    "target_id": para_element["parent_id"],
                    "relationship_type": RelationshipType.CONTAINED_BY.value,
                    "metadata": {
                        "confidence": 1.0
                    }
                }
                relationships.append(contained_by_relationship)

            elif isinstance(block, Table):
                # Process table
                table_elements, table_relationships = self._process_table(block, i, doc_id, current_parent, source_id)
                elements.extend(table_elements)
                relationships.extend(table_relationships)

        # Extract comments if enabled
        if self.extract_comments:
            try:
                comment_elements, comment_relationships = self._extract_comments(doc, doc_id, body_id, source_id)
                elements.extend(comment_elements)
                relationships.extend(comment_relationships)
            except Exception as e:
                logger.warning(f"Error extracting comments: {str(e)}")

        return elements, relationships

    def _process_paragraph(self, paragraph: Paragraph, index: int, doc_id: str, parent_id: str, source_id: str) -> \
            Optional[Dict[str, Any]]:
        """
        Process a paragraph element.

        Args:
            paragraph: The paragraph
            index: Element index
            doc_id: Document ID
            parent_id: Parent element ID
            source_id: Source identifier

        Returns:
            Paragraph element dictionary or None if empty
        """
        # Get text content
        text = paragraph.text.strip()

        # Skip empty paragraphs
        if not text:
            return None

        # Generate element ID
        element_id = self._generate_id("para_")

        # Get paragraph style
        style_name = paragraph.style.name if paragraph.style else "Normal"

        # Get alignment
        alignment = "left"  # Default
        if paragraph.paragraph_format:
            if paragraph.paragraph_format.alignment:
                alignment_value = paragraph.paragraph_format.alignment
                if alignment_value == 1:
                    alignment = "center"
                elif alignment_value == 2:
                    alignment = "right"
                elif alignment_value == 3:
                    alignment = "justified"

        # Create paragraph element
        element = {
            "element_id": element_id,
            "doc_id": doc_id,
            "element_type": ElementType.PARAGRAPH.value,
            "parent_id": parent_id,
            "content_preview": text[:self.max_content_preview] + (
                "..." if len(text) > self.max_content_preview else ""),
            "content_location": json.dumps({
                "source": source_id,
                "type": ElementType.PARAGRAPH.value,
                "index": index
            }),
            "content_hash": self._generate_hash(text),
            "metadata": {
                "style": style_name,
                "alignment": alignment,
                "index": index
            }
        }

        # Check for list formatting
        if paragraph._p.pPr and paragraph._p.pPr.numPr:
            # This is a list item
            element["element_type"] = ElementType.LIST_ITEM.value

            # Try to determine list type and level
            list_level = 0
            if paragraph._p.pPr.numPr.ilvl:
                list_level = int(paragraph._p.pPr.numPr.ilvl.val)

            element["metadata"]["list_level"] = list_level

            # List type is harder to determine reliably without full numbering definitions
            # For now, we'll make a guess based on the first character
            if text.startswith(('•', '○', '■', '●', '◦', '◆')):
                element["metadata"]["list_type"] = "unordered"
            elif text.strip()[0].isdigit() and text.strip()[1:3] in ('. ', '.) '):
                element["metadata"]["list_type"] = "ordered"
            else:
                element["metadata"]["list_type"] = "unknown"

        return element

    def _process_table(self, table: Table, index: int, doc_id: str, parent_id: str, source_id: str) -> tuple[
        List[Dict[str, Any]], List[Dict[str, Any]]]:
        """
        Process a table element and create relationships.

        Args:
            table: The table
            index: Element index
            doc_id: Document ID
            parent_id: Parent element ID
            source_id: Source identifier

        Returns:
            Tuple of (list of table-related elements, list of relationships)
        """
        elements = []
        relationships = []

        # Generate table element ID
        table_id = self._generate_id("table_")

        # Create table element
        table_preview = self._get_table_preview(table)
        table_element = {
            "element_id": table_id,
            "doc_id": doc_id,
            "element_type": ElementType.TABLE.value,
            "parent_id": parent_id,
            "content_preview": table_preview,  # This can be empty if we couldn't extract meaningful content
            "content_location": json.dumps({
                "source": source_id,
                "type": ElementType.TABLE.value,
                "index": index
            }),
            "content_hash": "",
            "metadata": {
                "rows": len(table.rows),
                "columns": len(table.columns),
                "index": index
            }
        }
        elements.append(table_element)

        # Create relationship from parent to table
        contains_relationship = {
            "relationship_id": self._generate_id("rel_"),
            "source_id": parent_id,
            "target_id": table_id,
            "relationship_type": RelationshipType.CONTAINS.value,
            "metadata": {
                "confidence": 1.0,
                "index": index
            }
        }
        relationships.append(contains_relationship)

        # Create inverse relationship
        contained_by_relationship = {
            "relationship_id": self._generate_id("rel_"),
            "source_id": table_id,
            "target_id": parent_id,
            "relationship_type": RelationshipType.CONTAINED_BY.value,
            "metadata": {
                "confidence": 1.0
            }
        }
        relationships.append(contained_by_relationship)

        # Process rows
        for row_idx, row in enumerate(table.rows):
            # Generate row element ID
            row_id = self._generate_id("row_")

            # Create row element
            row_element = {
                "element_id": row_id,
                "doc_id": doc_id,
                "element_type": ElementType.TABLE_ROW.value,
                "parent_id": table_id,
                "content_preview": f"Row {row_idx + 1}",
                "content_location": json.dumps({
                    "source": source_id,
                    "type": ElementType.TABLE_ROW.value,
                    "table_index": index,
                    "row": row_idx
                }),
                "content_hash": "",
                "metadata": {
                    "row": row_idx
                }
            }
            elements.append(row_element)

            # Create relationship from table to row
            contains_row_relationship = {
                "relationship_id": self._generate_id("rel_"),
                "source_id": table_id,
                "target_id": row_id,
                "relationship_type": RelationshipType.CONTAINS_TABLE_ROW.value,
                "metadata": {
                    "confidence": 1.0,
                    "row_index": row_idx
                }
            }
            relationships.append(contains_row_relationship)

            # Create inverse relationship
            row_contained_relationship = {
                "relationship_id": self._generate_id("rel_"),
                "source_id": row_id,
                "target_id": table_id,
                "relationship_type": RelationshipType.CONTAINED_BY.value,
                "metadata": {
                    "confidence": 1.0
                }
            }
            relationships.append(row_contained_relationship)

            # Process cells
            for col_idx, cell in enumerate(row.cells):
                # Generate cell element ID
                cell_id = self._generate_id("cell_")

                # Get cell content
                cell_text = " ".join(p.text for p in cell.paragraphs).strip()

                # Create cell element
                cell_element = {
                    "element_id": cell_id,
                    "doc_id": doc_id,
                    "element_type": ElementType.TABLE_CELL.value,
                    "parent_id": row_id,
                    "content_preview": cell_text[:self.max_content_preview] + (
                        "..." if len(cell_text) > self.max_content_preview else ""),
                    "content_location": json.dumps({
                        "source": source_id,
                        "type": ElementType.TABLE_CELL.value,
                        "table_index": index,
                        "row": row_idx,
                        "col": col_idx
                    }),
                    "content_hash": self._generate_hash(cell_text),
                    "metadata": {
                        "row": row_idx,
                        "col": col_idx,
                        "text": cell_text
                    }
                }

                # Check if this is a header cell (first row)
                if row_idx == 0:
                    cell_element["element_type"] = "table_header"

                elements.append(cell_element)

                # Create relationship from row to cell
                contains_cell_relationship = {
                    "relationship_id": self._generate_id("rel_"),
                    "source_id": row_id,
                    "target_id": cell_id,
                    "relationship_type": RelationshipType.CONTAINS_TABLE_CELL.value if cell_element[
                                                                                           "element_type"] == ElementType.TABLE_CELL.value else RelationshipType.CONTAINS_TABLE_HEADER.value,
                    "metadata": {
                        "confidence": 1.0,
                        "col_index": col_idx
                    }
                }
                relationships.append(contains_cell_relationship)

                # Create inverse relationship
                cell_contained_relationship = {
                    "relationship_id": self._generate_id("rel_"),
                    "source_id": cell_id,
                    "target_id": row_id,
                    "relationship_type": RelationshipType.CONTAINED_BY.value,
                    "metadata": {
                        "confidence": 1.0
                    }
                }
                relationships.append(cell_contained_relationship)

        return elements, relationships

    def _extract_headers_footers(self, doc: DocxDocument, doc_id: str, parent_id: str, source_id: str) -> tuple[
        List[Dict[str, Any]], List[Dict[str, Any]]]:
        """
        Extract headers and footers from document and create relationships.

        Args:
            doc: The DOCX document
            doc_id: Document ID
            parent_id: Parent element ID
            source_id: Source identifier

        Returns:
            Tuple of (list of header/footer elements, list of relationships)
        """
        elements = []
        relationships = []

        try:
            # Create headers container
            headers_id = self._generate_id("headers_")
            headers_element = {
                "element_id": headers_id,
                "doc_id": doc_id,
                "element_type": ElementType.HEADERS.value,
                "parent_id": parent_id,
                "content_preview": "Document headers",
                "content_location": json.dumps({
                    "source": source_id,
                    "type": ElementType.HEADERS.value
                }),
                "content_hash": "",
                "metadata": {}
            }
            elements.append(headers_element)

            # Create relationship from parent to headers container
            contains_headers_relationship = {
                "relationship_id": self._generate_id("rel_"),
                "source_id": parent_id,
                "target_id": headers_id,
                "relationship_type": RelationshipType.CONTAINS.value,
                "metadata": {
                    "confidence": 1.0
                }
            }
            relationships.append(contains_headers_relationship)

            # Create inverse relationship
            headers_contained_relationship = {
                "relationship_id": self._generate_id("rel_"),
                "source_id": headers_id,
                "target_id": parent_id,
                "relationship_type": RelationshipType.CONTAINED_BY.value,
                "metadata": {
                    "confidence": 1.0
                }
            }
            relationships.append(headers_contained_relationship)

            # Create footers container
            footers_id = self._generate_id("footers_")
            footers_element = {
                "element_id": footers_id,
                "doc_id": doc_id,
                "element_type": ElementType.FOOTERS.value,
                "parent_id": parent_id,
                "content_preview": "Document footers",
                "content_location": json.dumps({
                    "source": source_id,
                    "type": ElementType.FOOTERS.value
                }),
                "content_hash": "",
                "metadata": {}
            }
            elements.append(footers_element)

            # Create relationship from parent to footers container
            contains_footers_relationship = {
                "relationship_id": self._generate_id("rel_"),
                "source_id": parent_id,
                "target_id": footers_id,
                "relationship_type": RelationshipType.CONTAINS.value,
                "metadata": {
                    "confidence": 1.0
                }
            }
            relationships.append(contains_footers_relationship)

            # Create inverse relationship
            footers_contained_relationship = {
                "relationship_id": self._generate_id("rel_"),
                "source_id": footers_id,
                "target_id": parent_id,
                "relationship_type": RelationshipType.CONTAINED_BY.value,
                "metadata": {
                    "confidence": 1.0
                }
            }
            relationships.append(footers_contained_relationship)

            # Process sections
            for sect_idx, section in enumerate(doc.sections):
                # Process headers
                for header_type in ['first_page_header', 'header', 'even_page_header']:
                    header = getattr(section, header_type)
                    if header and header.is_linked_to_previous is False:
                        # Extract content
                        header_text = ""
                        for paragraph in header.paragraphs:
                            header_text += paragraph.text + "\n"

                        header_text = header_text.strip()
                        if header_text:
                            # Create header element
                            header_id = self._generate_id("header_")
                            header_element = {
                                "element_id": header_id,
                                "doc_id": doc_id,
                                "element_type": ElementType.PAGE_HEADER.value,
                                "parent_id": headers_id,
                                "content_preview": header_text[:self.max_content_preview] + (
                                    "..." if len(header_text) > self.max_content_preview else ""),
                                "content_location": json.dumps({
                                    "source": source_id,
                                    "type": ElementType.PAGE_HEADER.value,
                                    "section": sect_idx,
                                    "header_type": header_type
                                }),
                                "content_hash": self._generate_hash(header_text),
                                "metadata": {
                                    "section": sect_idx,
                                    "header_type": header_type.replace('_', ' '),
                                    "text": header_text
                                }
                            }
                            elements.append(header_element)

                            # Create relationship from headers container to header
                            contains_header_relationship = {
                                "relationship_id": self._generate_id("rel_"),
                                "source_id": headers_id,
                                "target_id": header_id,
                                "relationship_type": RelationshipType.CONTAINS.value,
                                "metadata": {
                                    "confidence": 1.0,
                                    "section": sect_idx,
                                    "type": header_type
                                }
                            }
                            relationships.append(contains_header_relationship)

                            # Create inverse relationship
                            header_contained_relationship = {
                                "relationship_id": self._generate_id("rel_"),
                                "source_id": header_id,
                                "target_id": headers_id,
                                "relationship_type": RelationshipType.CONTAINED_BY.value,
                                "metadata": {
                                    "confidence": 1.0
                                }
                            }
                            relationships.append(header_contained_relationship)

                # Process footers
                for footer_type in ['first_page_footer', 'footer', 'even_page_footer']:
                    footer = getattr(section, footer_type)
                    if footer and footer.is_linked_to_previous is False:
                        # Extract content
                        footer_text = ""
                        for paragraph in footer.paragraphs:
                            footer_text += paragraph.text + "\n"

                        footer_text = footer_text.strip()
                        if footer_text:
                            # Create footer element
                            footer_id = self._generate_id("footer_")
                            footer_element = {
                                "element_id": footer_id,
                                "doc_id": doc_id,
                                "element_type": ElementType.PAGE_FOOTER.value,
                                "parent_id": footers_id,
                                "content_preview": footer_text[:self.max_content_preview] + (
                                    "..." if len(footer_text) > self.max_content_preview else ""),
                                "content_location": json.dumps({
                                    "source": source_id,
                                    "type": ElementType.PAGE_FOOTER.value,
                                    "section": sect_idx,
                                    "footer_type": footer_type
                                }),
                                "content_hash": self._generate_hash(footer_text),
                                "metadata": {
                                    "section": sect_idx,
                                    "footer_type": footer_type.replace('_', ' '),
                                    "text": footer_text
                                }
                            }
                            elements.append(footer_element)

                            # Create relationship from footers container to footer
                            contains_footer_relationship = {
                                "relationship_id": self._generate_id("rel_"),
                                "source_id": footers_id,
                                "target_id": footer_id,
                                "relationship_type": RelationshipType.CONTAINS.value,
                                "metadata": {
                                    "confidence": 1.0,
                                    "section": sect_idx,
                                    "type": footer_type
                                }
                            }
                            relationships.append(contains_footer_relationship)

                            # Create inverse relationship
                            footer_contained_relationship = {
                                "relationship_id": self._generate_id("rel_"),
                                "source_id": footer_id,
                                "target_id": footers_id,
                                "relationship_type": RelationshipType.CONTAINED_BY.value,
                                "metadata": {
                                    "confidence": 1.0
                                }
                            }
                            relationships.append(footer_contained_relationship)
        except Exception as e:
            logger.warning(f"Error extracting headers/footers: {str(e)}")

        return elements, relationships

    def _extract_comments(self, doc: DocxDocument, doc_id: str, parent_id: str, source_id: str) -> tuple[
        List[Dict[str, Any]], List[Dict[str, Any]]]:
        """
        Extract comments from document and create relationships.

        Args:
            doc: The DOCX document
            doc_id: Document ID
            parent_id: Parent element ID
            source_id: Source identifier

        Returns:
            Tuple of (list of comment elements, list of relationships)
        """
        elements = []
        relationships = []

        try:
            # Create comments container
            comments_id = self._generate_id("comments_")
            comments_element = {
                "element_id": comments_id,
                "doc_id": doc_id,
                "element_type": ElementType.COMMENTS.value,
                "parent_id": parent_id,
                "content_preview": "Document comments",
                "content_location": json.dumps({
                    "source": source_id,
                    "type": ElementType.COMMENTS.value
                }),
                "content_hash": "",
                "metadata": {}
            }
            elements.append(comments_element)

            # Create relationship from parent to comments container
            contains_comments_relationship = {
                "relationship_id": self._generate_id("rel_"),
                "source_id": parent_id,
                "target_id": comments_id,
                "relationship_type": RelationshipType.CONTAINS.value,
                "metadata": {
                    "confidence": 1.0
                }
            }
            relationships.append(contains_comments_relationship)

            # Create inverse relationship
            comments_contained_relationship = {
                "relationship_id": self._generate_id("rel_"),
                "source_id": comments_id,
                "target_id": parent_id,
                "relationship_type": RelationshipType.CONTAINED_BY.value,
                "metadata": {
                    "confidence": 1.0
                }
            }
            relationships.append(comments_contained_relationship)

            # Extract comments
            # This is a bit tricky as python-docx doesn't have a direct API for comments
            # We need to access the XML directly

            # Get comments part if it exists
            if doc.part.package.parts:
                for rel_type, parts in doc.part.package.rels.items():
                    if 'comments' in rel_type.lower():
                        for rel_id, rel in parts.items():
                            if hasattr(rel, 'target_part') and rel.target_part:
                                # Process comments
                                comments_xml = rel.target_part.blob
                                if comments_xml:
                                    soup = BeautifulSoup(comments_xml, 'xml')
                                    for i, comment in enumerate(soup.find_all('comment')):
                                        # Extract comment data
                                        comment_id = comment.get('id', '')
                                        author = comment.get('author', 'Unknown')
                                        date = comment.get('date', '')
                                        text = comment.get_text().strip()

                                        # Create comment element
                                        comment_element_id = self._generate_id("comment_")
                                        comment_element = {
                                            "element_id": comment_element_id,
                                            "doc_id": doc_id,
                                            "element_type": ElementType.COMMENT.value,
                                            "parent_id": comments_id,
                                            "content_preview": text[:self.max_content_preview] + (
                                                "..." if len(text) > self.max_content_preview else ""),
                                            "content_location": json.dumps({
                                                "source": source_id,
                                                "type": ElementType.COMMENT.value,
                                                "comment_id": comment_id
                                            }),
                                            "content_hash": self._generate_hash(text),
                                            "metadata": {
                                                "comment_id": comment_id,
                                                "author": author,
                                                "date": date,
                                                "text": text,
                                                "index": i
                                            }
                                        }
                                        elements.append(comment_element)

                                        # Create relationship from comments container to comment
                                        contains_comment_relationship = {
                                            "relationship_id": self._generate_id("rel_"),
                                            "source_id": comments_id,
                                            "target_id": comment_element_id,
                                            "relationship_type": RelationshipType.CONTAINS.value,
                                            "metadata": {
                                                "confidence": 1.0,
                                                "index": i
                                            }
                                        }
                                        relationships.append(contains_comment_relationship)

                                        # Create inverse relationship
                                        comment_contained_relationship = {
                                            "relationship_id": self._generate_id("rel_"),
                                            "source_id": comment_element_id,
                                            "target_id": comments_id,
                                            "relationship_type": RelationshipType.CONTAINED_BY.value,
                                            "metadata": {
                                                "confidence": 1.0
                                            }
                                        }
                                        relationships.append(comment_contained_relationship)
        except Exception as e:
            logger.warning(f"Error extracting comments: {str(e)}")

        return elements, relationships

    def _extract_links(self, doc: DocxDocument, elements: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        Extract links from document.

        Args:
            doc: The DOCX document
            elements: Document elements

        Returns:
            List of extracted links
        """
        links = []

        try:
            # Extract hyperlinks
            # This is a bit tricky as python-docx doesn't have a direct API for hyperlinks
            # We need to access the XML directly

            # Iterate through relationships to find hyperlinks
            rels = doc.part.rels
            hyperlink_rels = {rel_id: rel.target_ref for rel_id, rel in rels.items()
                              if
                              rel.reltype == 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink'}

            # Iterate through all elements to find paragraphs
            for element in elements:
                element_id = element["element_id"]
                element_type = element["element_type"]

                # Only process text elements (paragraphs, headers, cells)
                if element_type not in ("paragraph", "header", "table_cell", "table_header", "list_item"):
                    continue

                # Get element content
                content_preview = element.get("content_preview", "")

                # Try to find hyperlinks in the XML for this element
                # This is a simplified approach that may not catch all hyperlinks
                for rel_id, target in hyperlink_rels.items():
                    # Check if this hyperlink's text appears in the element
                    # This is not perfect but a reasonable approximation
                    if target in content_preview:
                        # Create link
                        links.append({
                            "source_id": element_id,
                            "link_text": target,  # We don't know the exact text
                            "link_target": target,
                            "link_type": "hyperlink"
                        })
        except Exception as e:
            logger.warning(f"Error extracting links: {str(e)}")

        return links

    @staticmethod
    def _extract_styles_info(doc: DocxDocument) -> Dict[str, Any]:
        """
        Extract information about styles used in the document.

        Args:
            doc: The DOCX document

        Returns:
            Dictionary of style information
        """
        styles_info = {}

        try:
            # Get all styles in the document
            styles = doc.styles

            # Extract information about paragraph styles
            paragraph_styles = {}
            for style in styles:
                if style.type == 1:  # Paragraph style
                    paragraph_styles[style.name] = {
                        "style_id": style.style_id,
                        "based_on": style.base_style.name if style.base_style else None,
                        "builtin": not style.style_id.startswith('s')
                    }

            styles_info["paragraph_styles"] = paragraph_styles

            # Count usage of styles
            style_usage = {}
            for paragraph in doc.paragraphs:
                style_name = paragraph.style.name if paragraph.style else "Default"
                style_usage[style_name] = style_usage.get(style_name, 0) + 1

            styles_info["style_usage"] = style_usage

        except Exception as e:
            logger.warning(f"Error extracting style information: {str(e)}")

        return styles_info

    @staticmethod
    def _iter_block_items(doc: DocxDocument):
        """
        Iterate through all block items (paragraphs and tables) in document.

        Args:
            doc: The DOCX document

        Yields:
            Paragraph or Table objects
        """
        # Use XML to get all block items in order
        body = doc._body._body
        for child in body.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, doc._body)
            elif isinstance(child, CT_Tbl):
                yield Table(child, doc._body)

    @staticmethod
    def _estimate_page_count(doc: DocxDocument) -> int:
        """
        Estimate page count for document.

        Args:
            doc: The DOCX document

        Returns:
            Estimated page count
        """
        # This is just a rough estimation as actual page count depends on formatting
        # Assuming 250 words per page on average
        word_count = sum(len(paragraph.text.split()) for paragraph in doc.paragraphs)
        page_count = max(1, word_count // 250)
        return page_count

    @staticmethod
    def _count_words(doc: DocxDocument) -> int:
        """
        Count words in document.

        Args:
            doc: The DOCX document

        Returns:
            Word count
        """
        return sum(len(paragraph.text.split()) for paragraph in doc.paragraphs)

    def _get_table_preview(self, table: Table) -> str:
        """Generate a meaningful content preview for a table."""
        cell_texts = []

        # Try different methods to extract text from the table
        try:
            # First attempt: direct iteration (preferred when it works)
            for row in table.rows:
                for cell in row.cells:
                    text = " ".join(p.text for p in cell.paragraphs).strip()
                    if text:
                        cell_texts.append(text)
        except Exception:
            try:
                # Second attempt: index-based access
                for i in range(len(table.rows)):
                    row = table.rows[i]
                    for j in range(len(row.cells)):
                        cell = row.cells[j]
                        text = " ".join(p.text for p in cell.paragraphs).strip()
                        if text:
                            cell_texts.append(text)
            except Exception:
                # Third attempt: direct XML access as last resort
                try:
                    tbl_element = table._tbl
                    for text_elem in tbl_element.xpath('.//w:t'):
                        if text_elem.text and text_elem.text.strip():
                            cell_texts.append(text_elem.text.strip())
                except Exception:
                    # Give up and return empty string
                    return ""

        # If we have no text, return empty string
        if not cell_texts:
            return ""

        # Generate preview from collected texts
        preview = " | ".join(cell_texts[:5])
        if len(cell_texts) > 5:
            preview += "..."

        # Truncate if needed
        if len(preview) > self.max_content_preview:
            preview = preview[:self.max_content_preview] + "..."

        return preview
