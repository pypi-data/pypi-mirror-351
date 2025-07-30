from pathlib import Path
from typing import Iterable, List, Optional, Union, Tuple
import re
from bs4 import BeautifulSoup

import nltk
from nltk.corpus import stopwords
from nltk.stem import PorterStemmer

from sklearn.metrics.pairwise import cosine_similarity
from sklearn.feature_extraction.text import TfidfVectorizer

from stemmers.porter_ru import russian_porter_stem

from pdfminer.high_level import extract_text

import csv
from docx import Document as DocxDocument

import hashlib
import pickle
from datetime import datetime
import numpy as np
from scipy.sparse import save_npz, load_npz, vstack

import pdfplumber
import shutil

import warnings
warnings.filterwarnings("ignore")

import logging
logging.getLogger("pdfminer").setLevel(logging.ERROR)
logging.getLogger("PyPDF2").setLevel(logging.ERROR)

class Document:
    def __init__(self, path: Path, text: str, metadata: Optional[dict] = None):
        self.path = path
        self.text = text
        self.metadata = metadata or {}
        self.tokens: List[str] = []

    def __repr__(self) -> str:
        return f"<Document path={self.path} metadata={self.metadata}>"

class Loader:
    def __init__(self, path: Path, chunk_size: int = 500):
        self.path = path
        self.chunk_size = chunk_size

    def load(self) -> Iterable[List[Document]]:
        if not self.path.exists():
            raise FileNotFoundError(
                f"Указанный путь не существует: {self.path}\n"
                f"Рабочая директория: {Path.cwd()}\n"
                f"Абсолютный путь: {self.path.absolute()}"
            )
        
        files = [self.path] if self.path.is_file() else self.path.rglob('*')
        for file_path in files:
            if file_path.suffix.lower() in {'.txt', '.docx', '.csv', '.html', '.htm', '.pdf'}:
                yield from self._load_file(file_path)

    def _load_file(self, file_path: Path) -> Iterable[List[Document]]:
        handler_map = {
            '.txt': self._handle_text,
            '.docx': self._handle_docx,
            '.csv': self._handle_csv,
            '.pdf': self._handle_pdf,
            '.html': self._handle_html,
            '.htm': self._handle_html
        }
        
        handler = handler_map.get(file_path.suffix.lower())
        if handler:
            yield from handler(file_path)

    def _chunk_generator(self, elements: Iterable, file_path: Path, 
                        metadata_func: callable, label: str) -> Iterable[List[Document]]:
        chunk = []
        for idx, element in enumerate(elements, start=1):
            text = element.strip() if isinstance(element, str) else element
            if not text:
                continue
            
            metadata = metadata_func(idx)
            doc = Document(file_path, text, metadata)
            chunk.append(doc)
            
            if len(chunk) >= self.chunk_size:
                yield chunk
                chunk = []
        
        if chunk:
            yield chunk

    # Обработчики для разных форматов
    def _handle_text(self, file_path: Path) -> Iterable[List[Document]]:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            yield from self._chunk_generator(
                elements=(line.strip() for line in f),
                file_path=file_path,
                metadata_func=lambda idx: {'position': idx, 'label': 'line'},
                label='text'
            )

    def _handle_docx(self, file_path: Path) -> Iterable[List[Document]]:
        doc = DocxDocument(file_path)
        for para_num, para in enumerate(doc.paragraphs, start=1):
            text = para.text.strip()
            if text:
                yield [Document(
                    path=file_path,
                    text=text,
                    metadata={'position': para_num, 'label': 'paragraph', 'source': 'docx'}
                )]

    def _handle_pdf(self, file_path: Path) -> Iterable[List[Document]]:
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text(layout=False)  # Отключаем layout для более чистого текста
                    if not text:
                        continue
                    
                    # Обработка текста страницы
                    clean_text = re.sub(r'\s+', ' ', text).strip()
                    if clean_text:
                        yield [Document(
                            path=file_path,
                            text=clean_text,
                            metadata={
                                'page': page_num,
                                'label': 'page',
                                'source': 'pdf'
                            }
                        )]
        except Exception as e:
            print(f"PDF read error {file_path}: {str(e)}")

    def _handle_text(self, file_path: Path) -> Iterable[List[Document]]:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            current_block = []
            block_num = 1
            
            for line_num, line in enumerate(f, start=1):
                stripped = line.strip()
                if not stripped:
                    continue
                    
                current_block.append(stripped)
                
                # Собираем блоки по 3-5 строк
                if len(current_block) >= 3:
                    yield [Document(
                        path=file_path,
                        text=' '.join(current_block),
                        metadata={
                            'start_line': line_num - len(current_block) + 1,
                            'end_line': line_num,
                            'label': 'block',
                            'source': 'txt'
                        }
                    )]
                    current_block = []
                    block_num += 1
            
            if current_block:
                yield [Document(
                    path=file_path,
                    text=' '.join(current_block),
                    metadata={
                        'start_line': line_num - len(current_block) + 1,
                        'end_line': line_num,
                        'label': 'block',
                        'source': 'txt'
                    }
                )]

    def _handle_csv(self, file_path: Path) -> Iterable[List[Document]]:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            reader = csv.reader(f)
            header = next(reader, None)
            yield from self._chunk_generator(
                elements=(' '.join(row) for row in reader),
                file_path=file_path,
                metadata_func=lambda idx: {'position': idx+1, 'label': 'row', 'header': header},
                label='csv'
            )


    def _handle_html(self, file_path: Path) -> Iterable[List[Document]]:
        raw = file_path.read_text(encoding='utf-8', errors='ignore')
        yield [Document(path=file_path, text=raw, metadata={'suffix': file_path.suffix})]

class TextPipeline:
    def __init__(self, language: str = 'ru', use_stemming: bool = True, remove_stopwords: bool = True):
        self.language = language
        self.use_stemming = use_stemming
        self.remove_stopwords = remove_stopwords

        if remove_stopwords:
            try:
                if language == 'ru':
                    self.stopword_set = set(stopwords.words('russian'))
                elif language == 'en':
                    self.stopword_set = set(stopwords.words('english'))
                else:
                    self.stopword_set = set()
            except LookupError:
                print("Скачиваем ресурсы NLTK...")
                nltk.download('stopwords')
                nltk.download('punkt')
                
                if language == 'ru':
                    self.stopword_set = set(stopwords.words('russian'))
                elif language == 'en':
                    self.stopword_set = set(stopwords.words('english'))
                else:
                    self.stopword_set = set()
        else:
            self.stopword_set = set()

        if language == 'en' and use_stemming:
            self.stemmer_en = PorterStemmer()

    def strip_html(self, text: str) -> str:
        if '<' in text and '>' in text:
            return BeautifulSoup(text, 'html.parser').get_text(separator=' ')
        return text

    def lowercase(self, text: str) -> str:
        return text.lower()

    def remove_non_alphanumeric(self, text: str) -> str:
        # Удаляем все символы, кроме букв, цифр и пробелов
        text = re.sub(r'[^a-zа-яё0-9\s]', ' ', text.lower())
        # Заменяем множественные пробелы на один
        return re.sub(r'\s+', ' ', text).strip()

    def tokenize(self, text: str) -> List[str]:
        return text.split()

    def filter_stopwords(self, tokens: List[str]) -> List[str]:
        if not self.remove_stopwords:
            return tokens
        return [tok for tok in tokens if tok not in self.stopword_set]

    def stem_tokens(self, tokens: List[str]) -> List[str]:
        if not self.use_stemming:
            return tokens
        stemmed: List[str] = []
        for token in tokens:
            if self.language == 'ru':
                stemmed.append(russian_porter_stem(token))
            elif self.language == 'en':
                stemmed.append(self.stemmer_en.stem(token))
            else:
                stemmed.append(token)
        return stemmed

    def transform(self, text: str) -> List[str]:
        text = self.strip_html(text)
        text = self.lowercase(text)
        text = self.remove_non_alphanumeric(text)
        tokens = self.tokenize(text)
        tokens = self.filter_stopwords(tokens)
        tokens = self.stem_tokens(tokens)
        return tokens

def simple_tokenizer(x):
    return x.split()

class TfidfIndexer:
    def __init__(self):
        self.vectorizer = TfidfVectorizer(
            analyzer='word',
            tokenizer=simple_tokenizer,
            preprocessor=None,
            lowercase=False
        )
        self.matrix = None
        self.doc_index: List[Document] = []
        self._cache_version = "2.2"  # Обновляем версию

    def save(self, cache_dir: Path):
        try:
            cache_dir.mkdir(parents=True, exist_ok=True)
            if self.matrix is not None:
                save_npz(cache_dir / "matrix.npz", self.matrix)
            
            # Сохраняем векторайзер отдельно
            with open(cache_dir / "vectorizer.pkl", "wb") as f:
                pickle.dump(self.vectorizer, f)
            
            with open(cache_dir / "doc_index.pkl", "wb") as f:
                pickle.dump({
                    'version': self._cache_version,
                    'docs': self.doc_index,
                    'timestamp': datetime.now().timestamp()
                }, f)
        except Exception as e:
            print(f"Ошибка сохранения кэша: {str(e)}")
            if cache_dir.exists():
                shutil.rmtree(cache_dir)

    def load(self, cache_dir: Path) -> bool:
        try:
            self.matrix = load_npz(cache_dir / "matrix.npz")
            
            # Загружаем векторайзер
            with open(cache_dir / "vectorizer.pkl", "rb") as f:
                self.vectorizer = pickle.load(f)
            
            with open(cache_dir / "doc_index.pkl", "rb") as f:
                data = pickle.load(f)
                if data.get('version') != self._cache_version:
                    return False
                self.doc_index = data['docs']
            
            return True
        except Exception as e:
            print(f"Ошибка загрузки кэша: {str(e)}")
            return False
        
    def fit(self, token_texts: List[str], docs: List[Document]):
        """Обработка ВСЕХ документов за один раз"""
        self.doc_index = docs
        self.matrix = self.vectorizer.fit_transform(token_texts)

    def query(self, tokens: List[str], top_k: int = 5) -> List[Tuple[Document, float]]:
        query_text = ' '.join(tokens)
        vec = self.vectorizer.transform([query_text])
        sims = cosine_similarity(vec, self.matrix)[0]
        ranked = sorted(zip(self.doc_index, sims), key=lambda x: x[1], reverse=True)
        return ranked[:top_k] if top_k else ranked


class bkDetect:
    def __init__(self, input_path: Union[str, Path], language: str = 'ru', 
                 use_stemming: bool = True, remove_stopwords: bool = True, 
                 chunk_size: int = 500, use_cache: bool = True):
        self.path = Path(input_path)
        self.loader = Loader(self.path, chunk_size=chunk_size)
        self.pipeline = TextPipeline(language, use_stemming, remove_stopwords)
        self.indexer = TfidfIndexer()
        self.use_cache = use_cache
        self.cache_dir = self._get_cache_dir()
        if not self.path.exists():
            raise FileNotFoundError(f"Указанный путь не существует: {self.path}")

    def _get_cache_dir(self) -> Path:
        config_hash = hashlib.md5(f"{self.path}_{self.pipeline.language}_"
                                 f"{self.pipeline.use_stemming}_"
                                 f"{self.pipeline.remove_stopwords}".encode()).hexdigest()
        
        return Path(f".cache/{config_hash}")

    def _is_cache_valid(self) -> bool:
        if not self.use_cache or not self.cache_dir.exists():
            return False
        
        required_files = ['matrix.npz', 'doc_index.pkl', 'vectorizer.pkl']
        if not all((self.cache_dir / f).exists() for f in required_files):
            return False
        
        try:
            # Загружаем время создания кэша
            with open(self.cache_dir / "doc_index.pkl", "rb") as f:
                cache_data = pickle.load(f)
                cache_time = cache_data['timestamp']
            
            # Получаем список всех исходных файлов с абсолютными путями
            source_files = []
            if self.path.is_file():
                source_files = [self.path.absolute()]
            else:
                source_files = [f.absolute() for f in self.path.rglob('*') 
                                if f.is_file() and f.suffix.lower() in {'.txt', '.docx', '.csv', '.html', '.htm', '.pdf'}]
            
            if not source_files:
                return False
            
            # Получаем список файлов из кэша (абсолютные пути)
            with open(self.cache_dir / "doc_index.pkl", "rb") as f:
                cache_data = pickle.load(f)
                cache_docs = cache_data['docs']
                cache_files = {doc.path.absolute() for doc in cache_docs}
            
            # Проверяем соответствие файлов
            source_set = set(source_files)
            if source_set != cache_files:
                missing_files = source_set - cache_files
                extra_files = cache_files - source_set
                
                if missing_files:
                    print(f"Файлы отсутствуют в кэше: {[f.name for f in missing_files]}")
                if extra_files:
                    print(f"Лишние файлы в кэше: {[f.name for f in extra_files]}")
                return False
            
            # Проверяем время модификации
            newest_source = max(f.stat().st_mtime for f in source_files)
            if newest_source > cache_time:
                print(f"Обнаружены изменения в исходных файлах после создания кэша")
                return False
            
            return True
            
        except Exception as e:
            print(f"Ошибка проверки кэша: {str(e)}")
            return False

    def build_index(self) -> None:
        if self._is_cache_valid() and self.indexer.load(self.cache_dir):
            print("Используется кэшированный индекс")
            return

        print("Построение нового индекса...")
        all_docs = []
        token_texts = []
        
        for docs_chunk in self.loader.load():
            processed_chunk = []
            for doc in docs_chunk:
                doc = self._process_doc(doc)
                if doc.tokens:  # Пропускаем пустые документы
                    processed_chunk.append(doc)
            
            all_docs.extend(processed_chunk)
            token_texts.extend(' '.join(doc.tokens) for doc in processed_chunk)
        
        if not token_texts:
            print("Нет данных для индексации!")
            return
        
        self.indexer.fit(token_texts, all_docs)
        
        if self.use_cache:
            self.indexer.save(self.cache_dir)
            print(f"Индекс сохранен в кэш: {self.cache_dir}")

    def _process_doc(self, doc: Document) -> Document:
        doc.tokens = self.pipeline.transform(doc.text)
        return doc

    def find_sources(self, query_text: str, top_k: int = 5):
        tokens = self.pipeline.transform(query_text)
        raw = self.indexer.query(tokens, top_k=None)  # все
        best: Dict[Path, float] = {}
        for doc, score in raw:
            if score == 0.0:
                continue
            p = doc.path
            if p not in best or best[p] < score:
                best[p] = score

        sorted_paths = sorted(best.items(), key=lambda x: x[1], reverse=True)[:top_k]
        return [(path, best_score) for path, best_score in sorted_paths]


    def find_sources_from_file(self, file_path: Union[str, Path], top_k: int = 5) -> List[Tuple[Document, float]]:
        text = Path(file_path).read_text(encoding='utf-8', errors='ignore')
        return self.find_sources(text, top_k)


    def locate_source_positions(self, query_text: str, top_k: int = 5, max_positions_per_file: int = 2, snippet_len: int = 200):
        query_tokens = set(self.pipeline.transform(query_text))
        results = self.find_sources(query_text, top_k)
        result_paths = {path for path, _ in results}

        positions = []
        counts = {}

        for doc, score in self.indexer.query(self.pipeline.transform(query_text), top_k=None):
            if score < 0.1:  # Фильтр низких оценок
                continue
            if doc.path not in result_paths:
                continue
            if not query_tokens & set(doc.tokens):
                continue

            # ФОРМАТИРОВАНИЕ ПОЗИЦИЙ В ЗАВИСИМОСТИ ОТ ФОРМАТА
            meta = doc.metadata
            source_type = meta.get('source', 'unknown')
            pos_info = ""
            
            if source_type == 'pdf':
                pos_info = f"страница {meta.get('page', '?')}"
            elif source_type == 'docx':
                pos_info = f"line {meta.get('position', '?')}"
            elif source_type == 'txt':
                start = meta.get('start_line', '?')
                end = meta.get('end_line', '?')
                pos_info = f"строки {start}-{end}" if start != end else f"строка {start}"
            else:
                pos_info = f"позиция {meta.get('position', '?')}"

            snippet = doc.text
            if len(snippet) > snippet_len:
                snippet = snippet[:snippet_len].rstrip() + '…'
            
            file_key = doc.path
            if counts.get(file_key, 0) < max_positions_per_file:
                positions.append((doc.path, pos_info, snippet, score, source_type))
                counts[file_key] = counts.get(file_key, 0) + 1

        return positions