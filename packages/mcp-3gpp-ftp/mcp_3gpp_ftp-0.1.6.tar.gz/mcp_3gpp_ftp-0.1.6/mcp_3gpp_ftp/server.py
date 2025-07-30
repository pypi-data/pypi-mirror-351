"""
server.py

A FastMCP-based MCP Server exposing tools to crawl and list FTP directories and files under the 3GPP FTP site.

Usage:
    uv run server.py

This will start an MCP server on localhost:8000 (default) with the following tools:
    - list_directories(path: str) -> List[str]
    - crawl_ftp(path: str, depth: int, delay: float) -> List[str]
    - list_files(path: str) -> List[str]

Clients can introspect available tools and invoke them via the MCP protocol.
"""
from typing import List, Dict, Any, Optional
import fnmatch
import urllib.parse
import os
import time
import requests
from bs4 import BeautifulSoup
from io import BytesIO
import pandas as pd
from mcp.server.fastmcp import FastMCP
import zipfile
import base64
from docx import Document


# Instantiate the MCP server with a friendly name
mcp = FastMCP("MCP 3GPP FTP Explorer")

# Base URL of the 3GPP FTP directory
BASE_URL = "https://www.3gpp.org/ftp/"

proxies = {}

# ensure a cache directory exists next to this script
CACHE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "download_cache")
os.makedirs(CACHE_DIR, exist_ok=True)

@mcp.tool()
def list_directories_files(path: str = "", file_pattern: str = "*") -> List[str]:
    """
    List subdirectories and files under a given 3GPP FTP path, filtering files by wildcard.

    Sends an HTTP GET to BASE_URL + path, parses the HTML directory listing
    with BeautifulSoup, and returns:

      - All subdirectory names (no extension)  
      - Any files whose names match the given wildcard (fnmatch)

    Args:
        path (str):
            Relative FTP path (e.g. "Specifications/Rel.17/"), defaults to root (“”).
        file_pattern (str):
            Unix‐style wildcard (e.g. "*.xlsx", "*.zip"); only files matching this
            pattern are returned. Directories are always returned regardless of pattern.

    Returns:
        List[str]:
            Names of entries directly under `path`: all directories plus matching files.

    Raises:
        requests.HTTPError:
            On non-2xx HTTP responses.
        requests.RequestException:
            For network issues (timeouts, DNS errors, etc.).
        Exception:
            For parsing failures or other unexpected errors.
    """
    url = BASE_URL + path
    resp = requests.get(url, timeout=10, verify=False, proxies=proxies)
    resp.raise_for_status()

    soup = BeautifulSoup(resp.text, "html.parser")
    entries: List[str] = []
    for link in soup.find_all("a", href=True):
        href = link["href"]
        if not href.startswith(BASE_URL):
            continue

        name = os.path.basename(os.path.normpath(href.rstrip("/")))
        # determine if this is a directory (no dot in last segment)
        is_dir = "." not in name

        if is_dir:
            entries.append(name)
        else:
            # file: only add if it matches the pattern
            if fnmatch.fnmatch(name, file_pattern):
                entries.append(name)

    return entries
'''
def list_directories_files(path: str = "") -> List[str]:
    """
    Retrieve a list of subdirectory names and files at the given FTP path under the 3GPP FTP server.

    This function sends an HTTP GET request to the combined URL (BASE_URL + path), parses
    the HTML directory listing using BeautifulSoup, and extracts only the names of entries
    that represent subdirectories (i.e., links without a file extension in the last path segment).

    Args:
        path (str): The relative path (from the FTP root) to list. Defaults to the root ('').

    Returns:
        List[str]: A list of directory names (not full URLs) present under the specified path.

    Raises:
        requests.HTTPError: If the HTTP request returns a bad status code.
        requests.RequestException: For network-related errors (e.g., timeouts).
        Exception: For any other parsing or unexpected errors.

    Example:
        >>> list_directories_files("/")
        ["tsg_ran", "tsg_sa", "tsg_ct", "tsg_t"]
    """
    try:
        response = requests.get(BASE_URL + path, timeout=10, verify=False, proxies=proxies)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")
        dirs: List[str] = []
        for link in soup.find_all("a", href=True):
            href = link["href"]
            # Only consider links pointing into the BASE_URL tree, excluding files
            if href.startswith(BASE_URL):
                deepest_dir = os.path.basename(os.path.normpath(href.strip("/")))
                if deepest_dir not in BASE_URL + path:
                    dirs.append(deepest_dir) # Append only the deepest directory name
        return dirs
    except Exception as e:
        # Log and re-raise or return empty list by design
        print(f"Failed to list directories at {BASE_URL + path}: {e}")
        raise e
'''

@mcp.tool()
def list_directories(path: str = "") -> List[str]:
    """
    Retrieve a list of subdirectory names at the given FTP path under the 3GPP FTP server.

    This function sends an HTTP GET request to the combined URL (BASE_URL + path), parses
    the HTML directory listing using BeautifulSoup, and extracts only the names of entries
    that represent subdirectories (i.e., links without a file extension in the last path segment).

    Args:
        path (str): The relative path (from the FTP root) to list. Defaults to the root ('').

    Returns:
        List[str]: A list of directory names (not full URLs) present under the specified path.

    Raises:
        requests.HTTPError: If the HTTP request returns a bad status code.
        requests.RequestException: For network-related errors (e.g., timeouts).
        Exception: For any other parsing or unexpected errors.

    Example:
        >>> list_directories("/")
        ["tsg_ran", "Rel.18", "Rel.19"]
    """
    try:
        response = requests.get(BASE_URL + path, timeout=10, verify=False, proxies=proxies)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")
        dirs: List[str] = []
        for link in soup.find_all("a", href=True):
            href = link["href"]
            # Only consider links pointing into the BASE_URL tree, excluding files
            if href.startswith(BASE_URL) and "." not in href.rstrip("/").split('/')[-1]:
                deepest_dir = os.path.basename(os.path.normpath(href.strip("/")))
                if deepest_dir not in BASE_URL + path:
                    dirs.append(deepest_dir) # Append only the deepest directory name
        return dirs
    except Exception as e:
        # Log and re-raise or return empty list by design
        print(f"Failed to list directories at {BASE_URL + path}: {e}")
        raise e

#@mcp.tool()
def crawl_ftp(path: str = "", depth: int = 2, delay: float = 0.1) -> List[str]:
    """
    Recursively traverse the FTP directory tree up to a specified depth, collecting all discovered
    directory paths relative to the BASE_URL.

    This function calls `list_directories` at each level, sleeps briefly between requests to
    avoid overloading the server, and aggregates a flat list of all visited paths.

    Args:
        path (str): The starting relative path to crawl, defaults to root ('').
        depth (int): Maximum recursion depth; each level decrements this by 1.
        delay (float): Seconds to sleep between successive HTTP requests.

    Returns:
        List[str]: A list of all relative directory paths discovered (e.g., ['/Rel.17', '/Rel.17/TSG-RAN']).

    Raises:
        Exception: Propagates any exceptions from `list_directories` or network failures.

    Example:
        >>> crawl_ftp("/", depth=2)
        ['/tsg_ran', '/tsg_sa', '/Specs/']
    """
    all_paths: List[str] = []
    try:
        subdirs = list_directories(path)
        for sub in subdirs:
            new_path = f"{path.rstrip('/')}/{sub}"
            all_paths.append(new_path)
            if depth > 1:
                time.sleep(delay)
                all_paths.extend(crawl_ftp(new_path, depth - 1, delay))
        return all_paths
    except Exception as e:
        print(f"Error crawling FTP at path: {path}")
        raise e


#@mcp.tool()
def list_files(path: str = "") -> List[str]:
    """
    Retrieve a list of file names at the given FTP path under the 3GPP FTP server.

    Sends an HTTP GET to (BASE_URL + path), parses the HTML directory listing, and collects
    only those links whose last segment contains a file extension (i.e., contains a dot).

    Args:
        path (str): The relative path (from the FTP root) to list files. Defaults to root ('').

    Returns:
        List[str]: A list of file names (not full URLs) present under the specified path.

    Raises:
        requests.HTTPError: If the HTTP request returns a bad status code.
        requests.RequestException: For network-related errors (e.g., timeouts).
        Exception: For any other parsing or unexpected errors.

    Example:
        >>> list_files("/Specifications/Rel.17/TSG-RAN/")
        ['TSG_RAN_17_901.zip', 'TSG_RAN_17_901.pdf']
    """
    try:
        response = requests.get(BASE_URL + path, timeout=10, verify=False, proxies=proxies)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")
        files: List[str] = []
        for link in soup.find_all("a", href=True):
            href = link["href"]
            segment = href.rstrip("/").split('/')[-1]
            # Only include items with an extension
            if href.startswith(BASE_URL) and "." in segment:
                files.append(segment)
        return files
    except Exception as e:
        print(f"Failed to list files at {BASE_URL + path}: {e}")
        raise e


@mcp.tool()
def list_excel_columns(file_url: str) -> List[str]:
    """
    Retrieve all column names from the given .xlsx file.

    Args:
        file_url (str):
            The full URL (e.g. "https://…/MySpec.xlsx") or a local file path
            (you can use "file:///C:/path/to/file.xlsx") pointing to the XLSX file.

    Returns:
        List[str]: A list containing every column name in the first (active) sheet.

    Raises:
        HTTPError:           If downloading the file fails (non-200 status).
        ValueError:          If pandas cannot parse the content as Excel.
    """
    # 1. Download the file into memory
    resp = requests.get(file_url, timeout=10, verify=False, proxies=proxies)
    resp.raise_for_status()

    # 2. Load only the header row to get column names
    try:
        df = pd.read_excel(BytesIO(resp.content), nrows=0)
    except Exception as e:
        raise e

    # 3. Return the list of column names
    return list(df.columns)



@mcp.tool()
def filter_excel_columns_from_url(
    file_url: str,
    columns: List[str],
    filters: Optional[Dict[str, Any]] = None,
    match_all_keywords: bool = False
) -> List[Dict[str, Any]]:
    """
    Download an .xlsx from a URL, filter rows by substring (case-insensitive)
    on given column‐value pairs, then return only the specified columns.

    Args:
        file_url (str):
            Full HTTP(S) URL to an .xlsx file.
        columns (List[str]):
            List of column names to keep in the output.
        filters (Dict[str, Any], optional):
            Mapping of column name → value or list of values. Only rows
            where each column’s text contains (case-insensitive) the given
            value(s) are kept. If None or empty, no row filtering is applied.
        match_all_keywords (bool):
            If True, when a filter value is a list, require *all* keywords (AND).
            If False (default), require any keyword (OR).

    Returns:
        List[Dict[str, Any]]: A list of dicts (one per row) containing only the requested columns,
            and only rows matching the filters.

    Raises:
        ValueError:          • If file_url is not HTTP(S),  
                             • If parsing the Excel fails,  
                             • If any requested column or filter‐column is missing.  
        HTTPError:           If downloading the file fails (non-200 response).
    """
    if not file_url.lower().startswith(("http://", "https://")):
        raise ValueError(f"Only HTTP(S) URLs are supported, got: {file_url!r}")

    resp = requests.get(file_url, timeout=10, verify=False, proxies=proxies)
    resp.raise_for_status()
    data = BytesIO(resp.content)

    try:
        df = pd.read_excel(data, engine="openpyxl")
    except Exception as e:
        raise ValueError(f"Failed to parse Excel from URL: {e}")

    if filters:
        for fc, fv in filters.items():
            if fc not in df.columns:
                raise ValueError(f"Filter column not found in Excel: {fc}")
            series = df[fc].astype(str)
            if isinstance(fv, (list, tuple, set)):
                if match_all_keywords:
                    mask = True
                    for keyword in fv:
                        mask &= series.str.contains(str(keyword), case=False, na=False)
                else:
                    mask = False
                    for keyword in fv:
                        mask |= series.str.contains(str(keyword), case=False, na=False)
                df = df[mask]
            else:
                df = df[series.str.contains(str(fv), case=False, na=False)]

    missing = [c for c in columns if c not in df.columns]
    if missing:
        raise ValueError(f"Requested columns not found in Excel: {missing}")
    sub_df = df[columns]

    return sub_df.to_dict(orient="records")

@mcp.tool()
def download_and_extract(file_url: str) -> Dict[str, Any]:
    """
    Download a document (if not already cached), extract ZIPs to disk,
    and return:
      - output_path: the cache file or extraction folder
      - files: list of filenames present under that path
    """
    if not file_url.lower().startswith(("http://", "https://")):
        raise ValueError(f"Only HTTP(S) URLs are supported, got: {file_url!r}")

    filename = os.path.basename(file_url.split("?", 1)[0])
    cache_path = os.path.join(CACHE_DIR, filename)
    # Decide where we’ll dump files:
    if filename.lower().endswith(".zip"):
        output_path = os.path.join(CACHE_DIR, os.path.splitext(filename)[0])
    else:
        output_path = cache_path

    # If already there, just collect
    if os.path.exists(output_path):
        if os.path.isdir(output_path):
            files = os.listdir(output_path)
        else:
            files = [os.path.basename(output_path)]
        return {"output_path": output_path, "files": files}

    # Not there yet → download
    resp = requests.get(file_url, timeout=15, verify=False, proxies=proxies)
    resp.raise_for_status()
    with open(cache_path, "wb") as f:
        f.write(resp.content)

    # Now extract or keep
    if filename.lower().endswith(".zip"):
        os.makedirs(output_path, exist_ok=True)
        with zipfile.ZipFile(cache_path, "r") as z:
            for member in z.namelist():
                if member.endswith("/"):
                    continue
                target = os.path.join(output_path, member)
                os.makedirs(os.path.dirname(target), exist_ok=True)
                with z.open(member) as src, open(target, "wb") as dst:
                    dst.write(src.read())
        files = os.listdir(output_path)
    else:
        files = [filename]

    return {"output_path": output_path, "files": files}


@mcp.tool()
def read_word_doc(doc_path: str) -> Dict[str, Any]:
    """
    Read a .docx file from local disk and return its structure.

    Args:
        doc_path (str): Filesystem path to the .docx file.

    Returns:
        Dict[str, Any]: {
            "paragraphs": List[str],         # all paragraph texts in order
            "tables": List[List[List[str]]]  # a list of tables, each table is list of rows, each row is list of cell texts
        }

    Raises:
        FileNotFoundError: If `doc_path` does not exist.
        ValueError: If the file is not a valid .docx.
    """
    if not os.path.isfile(doc_path):
        raise FileNotFoundError(f"No such file: {doc_path}")

    try:
        doc = Document(doc_path)
    except Exception as e:
        raise ValueError(f"Failed to open as .docx: {e}")

    # Extract paragraphs
    paragraphs: List[str] = [p.text for p in doc.paragraphs]

    # Extract tables
    tables: List[List[List[str]]] = []
    for table in doc.tables:
        rows: List[List[str]] = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        tables.append(rows)

    return {
        "paragraphs": paragraphs,
        "tables": tables,
    }

def main():
    # Start the MCP server (default: localhost:8000)
    mcp.run(transport="stdio")

if __name__ == "__main__":
    # Start the MCP server (default: localhost:8000)
    main()
