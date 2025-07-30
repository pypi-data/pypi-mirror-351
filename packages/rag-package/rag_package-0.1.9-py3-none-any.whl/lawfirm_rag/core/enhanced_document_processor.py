"""
Enhanced Document Processing module for LawFirm-RAG.

Provides improved text extraction, processing, and document management
with support for vector databases and semantic search.
"""

import tempfile
import uuid
import re
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple, Union
import logging
from datetime import datetime
import hashlib
import json
import concurrent.futures
import threading
from functools import partial

# Text extraction libraries
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError as e:
    PYMUPDF_AVAILABLE = False
    logger.warning(f"PyMuPDF not available (common on Python 3.13): {e}")
    logger.info("PDF processing will use pdfplumber or PyPDF2 as fallback")

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Vector database
try:
    import chromadb
    CHROMADB_AVAILABLE = True
except ImportError:
    CHROMADB_AVAILABLE = False

logger = logging.getLogger(__name__)


class DocumentChunk:
    """Represents a chunk of a document with metadata."""
    
    def __init__(self, content: str, chunk_id: int, metadata: Dict[str, Any]):
        self.content = content
        self.chunk_id = chunk_id
        self.metadata = metadata
        self.embedding: Optional[List[float]] = None
        
    def to_dict(self) -> Dict[str, Any]:
        """Convert chunk to dictionary for storage."""
        return {
            "content": self.content,
            "chunk_id": self.chunk_id,
            "metadata": self.metadata,
            "embedding": self.embedding
        }


class ProcessedDocument:
    """Represents a processed document with chunks and metadata."""
    
    def __init__(self, filename: str, original_text: str, chunks: List[DocumentChunk], 
                 metadata: Dict[str, Any]):
        self.filename = filename
        self.original_text = original_text
        self.chunks = chunks
        self.metadata = metadata
        self.document_id = str(uuid.uuid4())
        
    def get_total_length(self) -> int:
        """Get total character length of all chunks."""
        return sum(len(chunk.content) for chunk in self.chunks)
        
    def to_dict(self) -> Dict[str, Any]:
        """Convert document to dictionary for storage."""
        return {
            "document_id": self.document_id,
            "filename": self.filename,
            "original_text": self.original_text,
            "chunks": [chunk.to_dict() for chunk in self.chunks],
            "metadata": self.metadata,
            "total_length": self.get_total_length()
        }


class EnhancedDocumentProcessor:
    """Enhanced document processor with improved text extraction and vector storage."""
    
    def __init__(self, 
                 temp_dir: Optional[str] = None,
                 chunk_size: int = 1000,
                 chunk_overlap: int = 200,
                 use_vector_db: bool = True,
                 vector_db_path: Optional[str] = None):
        """Initialize the enhanced document processor.
        
        Args:
            temp_dir: Directory for temporary file storage
            chunk_size: Size of text chunks for processing
            chunk_overlap: Overlap between chunks
            use_vector_db: Whether to use vector database storage
            vector_db_path: Path to vector database (None for in-memory)
        """
        self.temp_dir = Path(temp_dir) if temp_dir else Path(tempfile.gettempdir())
        
        # Ensure temp directory exists
        try:
            self.temp_dir.mkdir(parents=True, exist_ok=True)
            logger.info(f"Using temp directory: {self.temp_dir}")
        except Exception as e:
            logger.error(f"Failed to create temp directory {self.temp_dir}: {e}")
            # Fall back to system temp if custom temp dir fails
            self.temp_dir = Path(tempfile.gettempdir())
            logger.info(f"Falling back to system temp directory: {self.temp_dir}")
        
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.use_vector_db = use_vector_db
        
        # Collections store groups of documents (e.g., a case file)
        self.collections: Dict[str, Dict[str, Any]] = {}
        
        # Document storage - maps document_id to ProcessedDocument objects
        self.documents: Dict[str, ProcessedDocument] = {}
        
        # Initialize vector database if available and requested
        self.vector_db = None
        if use_vector_db and CHROMADB_AVAILABLE:
            try:
                db_path = vector_db_path or str(self.temp_dir / "vector_db")
                self.vector_db = chromadb.PersistentClient(path=db_path)
                logger.info(f"ChromaDB initialized successfully at: {db_path}")
            except Exception as e:
                logger.warning(f"Failed to initialize ChromaDB: {e}")
                self.vector_db = None
        
    def reset_vector_database(self):
        """Reset the vector database by deleting and recreating it."""
        if self.vector_db and CHROMADB_AVAILABLE:
            try:
                # Get current path - different methods for different ChromaDB versions
                db_path = None
                
                # Try different ways to get the database path
                if hasattr(self.vector_db, '_system') and hasattr(self.vector_db._system, '_path'):
                    db_path = self.vector_db._system._path
                elif hasattr(self.vector_db, '_settings') and hasattr(self.vector_db._settings, 'persist_directory'):
                    db_path = self.vector_db._settings.persist_directory
                elif hasattr(self.vector_db, '_path'):
                    db_path = self.vector_db._path
                
                # If we can't get the path, just recreate with default
                if not db_path:
                    logger.warning("Could not determine database path, using default temp directory")
                    db_path = str(self.temp_dir / "vector_db")
                
                # Close current connection
                del self.vector_db
                
                # Recreate database with fresh instance
                self.vector_db = chromadb.PersistentClient(path=db_path)
                logger.info(f"Vector database reset successfully at: {db_path}")
                
            except Exception as e:
                logger.error(f"Failed to reset vector database: {e}")
                # Fallback: try creating new database in temp directory
                try:
                    db_path = str(self.temp_dir / "vector_db_new")
                    self.vector_db = chromadb.PersistentClient(path=db_path)
                    logger.info(f"Created new vector database at: {db_path}")
                except Exception as e2:
                    logger.error(f"Failed to create new vector database: {e2}")
                    self.vector_db = None
        
    def create_collection(self, name: str, description: Optional[str] = None) -> str:
        """Create a new document collection.
        
        Args:
            name: Collection name (e.g., case number, matter name)
            description: Optional description of the collection
            
        Returns:
            Collection ID string
        """
        collection_id = str(uuid.uuid4())
        self.collections[collection_id] = {
            "name": name,
            "description": description,
            "documents": [],
            "created_at": datetime.now().isoformat(),
            "metadata": {}
        }
        
        # Create vector collection if available
        if self.vector_db:
            try:
                safe_name = re.sub(r'[^a-zA-Z0-9_-]', '_', name.lower())[:63]
                collection_name = f"collection_{safe_name}_{collection_id[:8]}"
                
                # Add complete metadata including the required topic field
                vector_collection = self.vector_db.create_collection(
                    name=collection_name,
                    metadata={
                        "collection_id": collection_id, 
                        "name": name,
                        "topic": description or name,  # Add the missing topic field
                        "description": description or "",
                        "created_at": datetime.now().isoformat()
                    }
                )
                logger.info(f"Created vector collection: {collection_name}")
                
                # Store collection reference
                self.collections[collection_id]["vector_collection_name"] = collection_name
                
            except Exception as e:
                logger.warning(f"Failed to create vector collection: {e}")
                # If it's a schema error, try resetting the database
                if "no such column" in str(e).lower():
                    logger.info("Attempting to reset vector database due to schema mismatch")
                    self.reset_vector_database()
                    # Try once more after reset
                    if self.vector_db:
                        try:
                            vector_collection = self.vector_db.create_collection(
                                name=collection_name,
                                metadata={
                                    "collection_id": collection_id, 
                                    "name": name,
                                    "topic": description or name,  # Add topic field here too
                                    "description": description or "",
                                    "created_at": datetime.now().isoformat()
                                }
                            )
                            logger.info(f"Created vector collection after reset: {collection_name}")
                            self.collections[collection_id]["vector_collection_name"] = collection_name
                        except Exception as e2:
                            logger.warning(f"Failed to create collection even after reset: {e2}")
                            logger.info("Continuing without vector collection - documents will still be processed")
                else:
                    logger.info("Continuing without vector collection - documents will still be processed")
        
        return collection_id
    
    def extract_text_advanced(self, file_path: Path) -> str:
        """Extract text using advanced methods with fallback options.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted and cleaned text content
        """
        ext = file_path.suffix.lower()
        text = ""
        
        if ext == ".pdf":
            text = self._extract_pdf_text(file_path)
        elif ext == ".docx":
            text = self._extract_docx_text(file_path)
        elif ext == ".txt":
            text = self._extract_txt_text(file_path)
        elif ext == ".json":
            text = self._extract_json_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
        
        # Clean and normalize the extracted text
        return self._clean_text(text)
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF using best available method."""
        text = ""
        
        # Try pdfplumber first (best for layout preservation)
        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n\n"
                logger.info(f"Extracted PDF text using pdfplumber: {len(text)} chars")
                return text
            except Exception as e:
                logger.warning(f"pdfplumber extraction failed: {e}")
        
        # Fallback to PyMuPDF
        if PYMUPDF_AVAILABLE:
            try:
                doc = fitz.open(file_path)
                for page_num in range(doc.page_count):
                    page = doc[page_num]
                    text += page.get_text() + "\n\n"
                doc.close()
                logger.info(f"Extracted PDF text using PyMuPDF: {len(text)} chars")
                return text
            except Exception as e:
                logger.warning(f"PyMuPDF extraction failed: {e}")
        
        # Last resort: PyPDF2
        if PYPDF2_AVAILABLE:
            try:
                with open(file_path, "rb") as file:
                    reader = PyPDF2.PdfReader(file)
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
                logger.info(f"Extracted PDF text using PyPDF2: {len(text)} chars")
                return text
            except Exception as e:
                logger.error(f"All PDF extraction methods failed: {e}")
                raise
        
        raise ImportError("No PDF processing libraries available")
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX processing")
        
        doc = docx.Document(file_path)
        paragraphs = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        
        return "\n\n".join(paragraphs)
    
    def _extract_txt_text(self, file_path: Path) -> str:
        """Extract text from TXT file with encoding detection."""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']
        
        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        
        raise ValueError(f"Could not decode text file with any common encoding")
    
    def _extract_json_text(self, file_path: Path) -> str:
        """Extract text from JSON file by converting to readable format."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            # Convert JSON to readable text format
            def json_to_text(obj, level=0):
                indent = "  " * level
                if isinstance(obj, dict):
                    text_parts = []
                    for key, value in obj.items():
                        if isinstance(value, (dict, list)):
                            text_parts.append(f"{indent}{key}:")
                            text_parts.append(json_to_text(value, level + 1))
                        else:
                            text_parts.append(f"{indent}{key}: {value}")
                    return "\n".join(text_parts)
                elif isinstance(obj, list):
                    text_parts = []
                    for i, item in enumerate(obj):
                        if isinstance(item, (dict, list)):
                            text_parts.append(f"{indent}Item {i + 1}:")
                            text_parts.append(json_to_text(item, level + 1))
                        else:
                            text_parts.append(f"{indent}- {item}")
                    return "\n".join(text_parts)
                else:
                    return f"{indent}{obj}"
            
            return json_to_text(data)
            
        except json.JSONDecodeError as e:
            logger.warning(f"Invalid JSON format in {file_path}: {e}")
            # Fallback: return raw content
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Failed to extract JSON text: {e}")
            return ""
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)  # Max 2 consecutive newlines
        text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs to single space
        text = re.sub(r'\r\n', '\n', text)  # Normalize line endings
        
        # Fix common OCR/extraction artifacts
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)  # Split camelCase
        text = re.sub(r'([.!?])([A-Z])', r'\1 \2', text)  # Space after punctuation
        
        # Remove page numbers and headers/footers (basic patterns)
        text = re.sub(r'^\d+\s*$', '', text, flags=re.MULTILINE)  # Standalone numbers
        text = re.sub(r'^Page \d+ of \d+.*$', '', text, flags=re.MULTILINE)  # Page headers
        
        return text.strip()
    
    def _chunk_text(self, text: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Split text into chunks with overlap."""
        if len(text) <= self.chunk_size:
            return [DocumentChunk(text, 0, metadata)]
        
        chunks = []
        start = 0
        chunk_id = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # If this isn't the last chunk, try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings within the last 200 characters
                search_start = max(end - 200, start)
                sentence_ends = []
                
                for pattern in [r'\.\s+', r'!\s+', r'\?\s+', r'\n\n']:
                    for match in re.finditer(pattern, text[search_start:end]):
                        sentence_ends.append(search_start + match.end())
                
                if sentence_ends:
                    end = max(sentence_ends)
            
            chunk_text = text[start:end].strip()
            if chunk_text:
                chunk_metadata = metadata.copy()
                chunk_metadata.update({
                    "chunk_id": chunk_id,
                    "start_char": start,
                    "end_char": end,
                    "chunk_length": len(chunk_text)
                })
                
                chunks.append(DocumentChunk(chunk_text, chunk_id, chunk_metadata))
                chunk_id += 1
            
            # Move start position with overlap
            start = end - self.chunk_overlap
            if start >= len(text):
                break
        
        return chunks
    
    def _extract_metadata(self, text: str, filename: str) -> Dict[str, Any]:
        """Extract metadata from document text."""
        metadata = {
            "filename": filename,
            "text_length": len(text),
            "processed_at": datetime.now().isoformat(),
            "word_count": len(text.split()),
            "document_hash": hashlib.md5(text.encode()).hexdigest()
        }
        
        # Basic pattern matching for legal document types
        text_lower = text.lower()
        
        # Document type detection
        if any(term in text_lower for term in ["complaint", "petition", "plaintiff", "defendant"]):
            metadata["document_type"] = "pleading"
        elif any(term in text_lower for term in ["contract", "agreement", "whereas"]):
            metadata["document_type"] = "contract"
        elif any(term in text_lower for term in ["deposition", "transcript", "sworn"]):
            metadata["document_type"] = "deposition"
        elif "intake" in text_lower or "questionnaire" in text_lower:
            metadata["document_type"] = "intake"
        else:
            metadata["document_type"] = "unknown"
        
        # Extract parties (basic pattern)
        parties = []
        party_patterns = [
            r'(?:plaintiff|defendant|petitioner|respondent)(?:\s+is\s+|\s*[:\-]\s*)([A-Z][a-z]+(?: [A-Z][a-z]+)*)',
            r'([A-Z][a-z]+(?: [A-Z][a-z]+)*)\s+(?:v\.|vs\.|versus)\s+([A-Z][a-z]+(?: [A-Z][a-z]+)*)'
        ]
        
        for pattern in party_patterns:
            matches = re.findall(pattern, text, re.MULTILINE)
            for match in matches:
                if isinstance(match, tuple):
                    parties.extend(match)
                else:
                    parties.append(match)
        
        if parties:
            metadata["parties"] = list(set(parties))
        
        # Extract dates (basic pattern)
        date_pattern = r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b'
        dates = re.findall(date_pattern, text)
        if dates:
            metadata["dates"] = dates[:5]  # Limit to first 5 dates found
        
        return metadata
    
    def process_document(self, file_path: Path, collection_id: str) -> ProcessedDocument:
        """Process a single document and add it to a collection.
        
        Args:
            file_path: Path to the document file
            collection_id: ID of the collection to add the document to
            
        Returns:
            ProcessedDocument object
        """
        if collection_id not in self.collections:
            raise ValueError(f"Collection {collection_id} does not exist")
        
        # Extract text
        text = self.extract_text_advanced(file_path)
        
        # Extract metadata
        metadata = self._extract_metadata(text, file_path.name)
        metadata["collection_id"] = collection_id
        
        # Create chunks
        chunks = self._chunk_text(text, metadata)
        
        # Create processed document
        processed_doc = ProcessedDocument(
            filename=file_path.name,
            original_text=text,
            chunks=chunks,
            metadata=metadata
        )
        
        # Store document for retrieval
        self.documents[processed_doc.document_id] = processed_doc
        
        # Add to collection
        self.collections[collection_id]["documents"].append(processed_doc.document_id)
        
        return processed_doc
    
    def _fast_extract_pdf_text(self, file_path: Path) -> str:
        """Fast PDF text extraction using the best available method."""
        # Use PyMuPDF first for speed if available
        if PYMUPDF_AVAILABLE:
            try:
                doc = fitz.open(file_path)
                text_parts = []
                for page_num in range(doc.page_count):
                    page = doc[page_num]
                    text_parts.append(page.get_text())
                doc.close()
                return "\n\n".join(text_parts)
            except Exception as e:
                logger.warning(f"PyMuPDF fast extraction failed: {e}")
        
        # Fallback to pdfplumber
        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(file_path) as pdf:
                    text_parts = []
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                return "\n\n".join(text_parts)
            except Exception as e:
                logger.warning(f"pdfplumber fast extraction failed: {e}")
        
        # Last resort: use the original method
        return self._extract_pdf_text(file_path)
    
    def _fast_chunk_text(self, text: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Fast text chunking with simplified boundary detection."""
        if len(text) <= self.chunk_size:
            return [DocumentChunk(text, 0, metadata)]
        
        chunks = []
        start = 0
        chunk_id = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # Simple boundary detection: look for nearest space
            if end < len(text):
                # Find the last space within 100 chars of the boundary
                search_start = max(end - 100, start + 100)  # Don't make chunks too small
                space_pos = text.rfind(' ', search_start, end)
                if space_pos > search_start:
                    end = space_pos
            
            chunk_text = text[start:end].strip()
            if chunk_text:
                chunk_metadata = metadata.copy()
                chunk_metadata.update({
                    "chunk_id": chunk_id,
                    "start_char": start,
                    "end_char": end,
                    "chunk_length": len(chunk_text)
                })
                
                chunks.append(DocumentChunk(chunk_text, chunk_id, chunk_metadata))
                chunk_id += 1
            
            # Move start position with overlap
            start = end - self.chunk_overlap
            if start >= len(text):
                break
        
        return chunks
    
    def _fast_extract_metadata(self, text: str, filename: str) -> Dict[str, Any]:
        """Fast metadata extraction with essential info only."""
        return {
            "filename": filename,
            "text_length": len(text),
            "processed_at": datetime.now().isoformat(),
            "word_count": len(text.split()),
            "document_hash": hashlib.md5(text.encode()).hexdigest(),
            "document_type": "legal_document"  # Simplified - detailed analysis can be done later
        }
    
    def _process_single_file_fast(self, file_data: Tuple[Any, str]) -> Tuple[Optional[ProcessedDocument], Optional[str]]:
        """Process a single file quickly - designed for parallel execution."""
        file, collection_id = file_data
        temp_path = None
        
        try:
            # Get filename safely
            filename = getattr(file, 'filename', 'unknown')
            if not filename or filename == 'unknown':
                filename = f"document_{threading.current_thread().ident}.pdf"
            
            # Sanitize filename
            safe_filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
            
            # Create unique temp file path
            temp_path = self.temp_dir / f"{uuid.uuid4()}_{safe_filename}"
            
            # Handle file content
            if hasattr(file, 'read'):
                content = file.read()
                if hasattr(file, 'seek'):
                    file.seek(0)
            else:
                content = file
            
            # Ensure content is bytes
            if isinstance(content, str):
                content = content.encode('utf-8')
            elif not isinstance(content, bytes):
                raise ValueError(f"Invalid file content type: {type(content)}")
            
            if not content:
                raise ValueError("File appears to be empty")
            
            # Write temp file
            with open(temp_path, "wb") as temp_file:
                temp_file.write(content)
            
            # Fast text extraction
            ext = temp_path.suffix.lower()
            if ext == ".pdf":
                text = self._fast_extract_pdf_text(temp_path)
            elif ext == ".docx":
                text = self._extract_docx_text(temp_path)
            elif ext == ".txt":
                text = self._extract_txt_text(temp_path)
            elif ext == ".json":
                text = self._extract_json_text(temp_path)
            else:
                raise ValueError(f"Unsupported file format: {ext}")
            
            # Clean text
            text = self._clean_text(text)
            
            # Fast metadata extraction
            metadata = self._fast_extract_metadata(text, temp_path.name)
            metadata["collection_id"] = collection_id
            
            # Fast chunking
            chunks = self._fast_chunk_text(text, metadata)
            
            # Create processed document
            processed_doc = ProcessedDocument(
                filename=temp_path.name,
                original_text=text,
                chunks=chunks,
                metadata=metadata
            )
            
            return processed_doc, None
            
        except Exception as e:
            error_msg = f"Error processing {getattr(file, 'filename', 'unknown')}: {e}"
            logger.error(error_msg)
            return None, error_msg
            
        finally:
            # Clean up temp file
            if temp_path and temp_path.exists():
                try:
                    temp_path.unlink()
                except Exception as cleanup_error:
                    logger.warning(f"Failed to clean up temp file {temp_path}: {cleanup_error}")
    
    def process_uploaded_files_fast(self, files: List[Any], collection_id: str, max_workers: int = 4) -> Dict[str, Any]:
        """Process multiple uploaded files in parallel for much faster processing.
        
        Args:
            files: List of uploaded file objects
            collection_id: ID of the collection to add documents to
            max_workers: Maximum number of parallel workers (default: 4)
            
        Returns:
            Processing results dictionary
        """
        if collection_id not in self.collections:
            raise ValueError(f"Collection {collection_id} does not exist")
        
        # Prepare file data for parallel processing
        file_data = [(file, collection_id) for file in files]
        
        processed_documents = []
        errors = []
        
        # Process files in parallel
        with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
            # Submit all files for processing
            future_to_file = {
                executor.submit(self._process_single_file_fast, data): data[0] 
                for data in file_data
            }
            
            # Collect results as they complete
            for future in concurrent.futures.as_completed(future_to_file):
                processed_doc, error = future.result()
                
                if processed_doc:
                    # Store document and add to collection
                    self.documents[processed_doc.document_id] = processed_doc
                    self.collections[collection_id]["documents"].append(processed_doc.document_id)
                    processed_documents.append(processed_doc)
                    logger.info(f"Successfully processed: {processed_doc.filename}")
                
                if error:
                    errors.append(error)
        
        return {
            "collection_id": collection_id,
            "processed_documents": len(processed_documents),
            "total_chunks": sum(len(doc.chunks) for doc in processed_documents),
            "total_text_length": sum(doc.get_total_length() for doc in processed_documents),
            "documents": [doc.to_dict() for doc in processed_documents],
            "errors": errors
        }
    
    def process_uploaded_files(self, files: List[Any], collection_id: str, use_fast_mode: bool = True) -> Dict[str, Any]:
        """Process multiple uploaded files and add them to a collection.
        
        Args:
            files: List of uploaded file objects
            collection_id: ID of the collection to add documents to
            use_fast_mode: Whether to use parallel processing (default: True)
            
        Returns:
            Processing results dictionary
        """
        if use_fast_mode:
            return self.process_uploaded_files_fast(files, collection_id)
        
        # Original sequential processing for compatibility
        if collection_id not in self.collections:
            raise ValueError(f"Collection {collection_id} does not exist")
        
        processed_documents = []
        errors = []
        
        for file in files:
            temp_path = None
            try:
                # Get filename safely
                filename = getattr(file, 'filename', 'unknown')
                if not filename or filename == 'unknown':
                    filename = f"document_{len(processed_documents)}.pdf"
                
                # Sanitize filename for filesystem safety
                safe_filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
                
                # Create unique temp file path
                temp_path = self.temp_dir / f"{uuid.uuid4()}_{safe_filename}"
                
                # Handle different file object types
                if hasattr(file, 'read'):
                    content = file.read()
                    if hasattr(file, 'seek'):
                        file.seek(0)
                else:
                    content = file
                
                # Ensure content is bytes
                if isinstance(content, str):
                    content = content.encode('utf-8')
                elif not isinstance(content, bytes):
                    raise ValueError(f"Invalid file content type: {type(content)}")
                
                # Verify we have content
                if not content:
                    raise ValueError("File appears to be empty")
                
                # Write file with explicit error handling
                logger.debug(f"Writing temp file: {temp_path}")
                with open(temp_path, "wb") as temp_file:
                    temp_file.write(content)
                
                # Verify file was written successfully
                if not temp_path.exists():
                    raise FileNotFoundError(f"Failed to create temp file: {temp_path}")
                
                file_size = temp_path.stat().st_size
                if file_size == 0:
                    raise ValueError(f"Temp file is empty: {temp_path}")
                
                logger.debug(f"Temp file created successfully: {temp_path} ({file_size} bytes)")
                
                # Process document
                processed_doc = self.process_document(temp_path, collection_id)
                processed_documents.append(processed_doc)
                
                logger.info(f"Successfully processed: {processed_doc.filename}")
                
            except Exception as e:
                error_msg = f"Error processing {getattr(file, 'filename', 'unknown')}: {e}"
                logger.error(error_msg)
                errors.append(error_msg)
                
            finally:
                # Clean up temp file regardless of success/failure
                if temp_path and temp_path.exists():
                    try:
                        temp_path.unlink()
                        logger.debug(f"Cleaned up temp file: {temp_path}")
                    except Exception as cleanup_error:
                        logger.warning(f"Failed to clean up temp file {temp_path}: {cleanup_error}")
        
        return {
            "collection_id": collection_id,
            "processed_documents": len(processed_documents),
            "total_chunks": sum(len(doc.chunks) for doc in processed_documents),
            "total_text_length": sum(doc.get_total_length() for doc in processed_documents),
            "documents": [doc.to_dict() for doc in processed_documents],
            "errors": errors
        }
    
    def get_collection_info(self, collection_id: str) -> Dict[str, Any]:
        """Get information about a collection.
        
        Args:
            collection_id: Collection ID
            
        Returns:
            Collection information dictionary
        """
        if collection_id not in self.collections:
            raise ValueError(f"Collection {collection_id} does not exist")
        
        collection = self.collections[collection_id].copy()
        
        # Add actual document data to collection info
        collection["documents"] = [
            self.documents[doc_id].to_dict() 
            for doc_id in collection["documents"] 
            if doc_id in self.documents
        ]
        
        return collection
    
    def get_collection_documents(self, collection_id: str) -> List[ProcessedDocument]:
        """Get all ProcessedDocument objects for a collection.
        
        Args:
            collection_id: Collection ID
            
        Returns:
            List of ProcessedDocument objects
        """
        if collection_id not in self.collections:
            raise ValueError(f"Collection {collection_id} does not exist")
        
        return [
            self.documents[doc_id] 
            for doc_id in self.collections[collection_id]["documents"] 
            if doc_id in self.documents
        ]
    
    def get_combined_text(self, collection_id: str) -> str:
        """Get combined text from all documents in a collection.
        
        Args:
            collection_id: Collection ID
            
        Returns:
            Combined text from all documents
        """
        documents = self.get_collection_documents(collection_id)
        return "\n\n".join([doc.original_text for doc in documents])
    
    def get_document_by_id(self, document_id: str) -> Optional[ProcessedDocument]:
        """Get a specific document by its ID.
        
        Args:
            document_id: Document ID
            
        Returns:
            ProcessedDocument object or None if not found
        """
        return self.documents.get(document_id)
    
    def list_collections(self) -> List[Dict[str, Any]]:
        """List all collections.
        
        Returns:
            List of collection information dictionaries
        """
        return [
            {"id": cid, **info} 
            for cid, info in self.collections.items()
        ]
    
    def cleanup_collection(self, collection_id: str) -> None:
        """Clean up a collection and its associated data.
        
        Args:
            collection_id: Collection ID to clean up
        """
        if collection_id not in self.collections:
            return
        
        # Clean up documents from memory
        document_ids = self.collections[collection_id].get("documents", [])
        for doc_id in document_ids:
            if doc_id in self.documents:
                del self.documents[doc_id]
        
        # Remove from vector database if available
        if self.vector_db:
            try:
                collection_name = f"collection_{collection_id[:8]}"
                self.vector_db.delete_collection(name=collection_name)
            except Exception as e:
                logger.warning(f"Could not delete vector collection: {e}")
        
        # Remove from memory
        del self.collections[collection_id]
        logger.info(f"Cleaned up collection: {collection_id}") 